from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from typing import Dict, List, Optional
import re


def _set_cell_shading(cell, color="E6E6E6"):
    """Apply gray background shading to a table cell."""
    shading = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
    cell._tc.get_or_add_tcPr().append(shading)


def _add_bold_run(paragraph, text, size=None, font_name=None):
    """Add a bold run to a paragraph."""
    run = paragraph.add_run(text)
    run.bold = True
    if size:
        run.font.size = size
    if font_name:
        run.font.name = font_name
    return run


def _add_run(paragraph, text, size=None, font_name=None, bold=False):
    """Add a run to a paragraph."""
    run = paragraph.add_run(text)
    run.bold = bold
    if size:
        run.font.size = size
    if font_name:
        run.font.name = font_name
    return run


def _extract_value_and_unit(name: str, raw_value: str):
    """
    Parse a spec value to extract:
    - display_value: what goes into column 2 (with ≥/≤ for constraints)
    - unit: what goes into column 3

    Examples:
      "не менее 16 Гб"      → ("≥16", "Гб")
      "не менее 2.4 ГГц"    → ("≥2,4", "ГГц")
      "не более 1.8 кг"     → ("≤1.8", "кг")
      "не менее 1920x1080"  → ("не менее 1920x1080", "")
      "SSD"                 → ("SSD", "")
      "120 Вт"              → ("≥120", "Вт") — if name is constraint-like
    """
    val = raw_value.strip()
    unit = ""

    # Known units (order matters — longer first to avoid partial matches)
    known_units = [
        'Кд/м²', 'Мбит/с', 'дюймов', 'пикселей', 'метра', 'метров',
        'ГГц', 'МГц', 'ГБ', 'Гб', 'МБ', 'MB', 'GHz', 'MHz',
        'Вт', 'кг', 'Шт', 'шт', 'dpi', 'мм', 'см', 'м',
    ]

    # Pattern 1: "не менее X unit" or "не более X unit"
    for constraint, symbol in [('не менее', '≥'), ('не более', '≤')]:
        if val.lower().startswith(constraint):
            rest = val[len(constraint):].strip()
            # Try to extract number + optional unit
            m = re.match(r'^(\d+[.,]?\d*)\s*(.*)?$', rest)
            if m:
                num = m.group(1).replace('.', ',')
                remainder = (m.group(2) or '').strip()
                # Check if remainder is a known unit
                matched_unit = ""
                for u in known_units:
                    if remainder.lower().startswith(u.lower()):
                        matched_unit = u
                        break
                if not matched_unit and remainder:
                    matched_unit = remainder
                return f"{symbol}{num}", matched_unit
            else:
                # "не менее 1920x1080 (FHD)" — keep as-is
                return val, ""

    # Pattern 2: Pure "NUMBER UNIT" — apply ≥ only if the name suggests a constraint
    m = re.match(r'^(\d+[.,]?\d*)\s+(.+)$', val)
    if m:
        num = m.group(1)
        remainder = m.group(2).strip()
        name_lower = name.lower()
        is_constraint = any(kw in name_lower for kw in [
            'количество', 'объем', 'объём', 'частота', 'мощность',
            'кэш', 'разъем', 'разъём', 'диагональ', 'яркость',
            'разрешение', 'длина', 'вес', 'число', 'максимальн',
            'тактовая', 'слот',
        ])
        if is_constraint:
            matched_unit = ""
            for u in known_units:
                if remainder.lower().startswith(u.lower()):
                    matched_unit = u
                    break
            if not matched_unit:
                matched_unit = remainder
            return f"≥{num}", matched_unit

    # Pattern 3: Bare number check for constraint-like names
    m = re.match(r'^(\d+[.,]?\d*)$', val)
    if m:
        name_lower = name.lower()
        is_constraint = any(kw in name_lower for kw in [
            'количество', 'объем', 'объём', 'частота', 'мощность',
            'кэш', 'разъем', 'разъём', 'число', 'слот',
        ])
        if is_constraint:
            return f"≥{m.group(1)}", "Шт."

    # Default: return as-is
    return val, unit


def _add_numbered_paragraph(doc, number: str, text: str, bold_number=True, bold_text=False, indent_cm=0):
    """Add a numbered paragraph like '1.1. Some text'."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if indent_cm > 0:
        p.paragraph_format.left_indent = Cm(indent_cm)

    if bold_number:
        _add_bold_run(p, f"{number}\t", size=Pt(11), font_name='Times New Roman')
    else:
        _add_run(p, f"{number}\t", size=Pt(11), font_name='Times New Roman')

    _add_run(p, text, size=Pt(11), font_name='Times New Roman', bold=bold_text)
    return p


def generate_tz_document(
    template_path: str,
    output_path: str,
    products: List[Dict],
    metadata: Optional[Dict] = None
):
    """
    Generates a formal TZ (Technical Specification) document.

    metadata structure:
    {
        "product_title": "системные блоки",
        "zakazchik": "ООО «Компания»",
        "quantity": 30,
        "quantity_text": "тридцать"
    }
    """
    doc = Document(template_path)

    # Default metadata
    meta = metadata or {}
    product_title = meta.get('product_title', 'оборудование')
    zakazchik = meta.get('zakazchik', '')
    quantity = meta.get('quantity', 1)
    quantity_text = meta.get('quantity_text', '')

    # Remove placeholder if exists
    for p in doc.paragraphs:
        if "{{PRODUCT_TABLES_PLACEHOLDER}}" in p.text:
            p.text = ""
            break

    # ── Section 1: Наименование, Заказчик, Исполнитель ──
    _add_numbered_paragraph(doc, "1.", "Наименование, Заказчик, Исполнитель, сроки и адрес поставки", bold_text=True)
    _add_numbered_paragraph(doc, "1.1.", f"Наименование объекта поставки: {product_title}.", indent_cm=0.5)
    _add_numbered_paragraph(doc, "1.2.", f"Заказчик: {zakazchik}", indent_cm=0.5)
    _add_numbered_paragraph(doc, "1.3.", "Исполнитель: определяется по результатам закупочных процедур.", indent_cm=0.5)

    # ── Section 2: Требования к поставке ──
    _add_numbered_paragraph(doc, "2.", "Требования к поставке Товара", bold_text=True)

    qty_display = f"{quantity}"
    if quantity_text:
        qty_display = f"{quantity} ({quantity_text})"
    _add_numbered_paragraph(doc, "2.1.", f"Требования к количеству поставляемого Товара: {qty_display} штук.", indent_cm=0.5)
    _add_numbered_paragraph(doc, "2.2.", "Требования к качеству поставляемого Товару:", indent_cm=0.5)

    # ── Product Tables ──
    for product_data in products:
        _add_product_table(doc, product_data)

    doc.save(output_path)
    print(f"Document saved to {output_path}")
    return output_path


def _add_product_table(document, product_data: Dict):
    """Add a single product's specification table with centered title."""
    p_name = product_data.get('product_name') or product_data.get('name') or "Товар"

    # Centered product title above the table
    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(8)
    title_p.paragraph_format.space_after = Pt(4)
    _add_bold_run(title_p, p_name, size=Pt(11), font_name='Times New Roman')

    # Create 3-column table
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    # Set column widths (approximate)
    for row in table.rows:
        row.cells[0].width = Cm(6.5)
        row.cells[1].width = Cm(6.5)
        row.cells[2].width = Cm(4.0)

    # Header row
    hdr = table.rows[0].cells
    headers = ['Наименование характеристики', 'Значение характеристики', 'Единица измерения характеристики']
    for i, (cell, text) in enumerate(zip(hdr, headers)):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        _set_cell_shading(cell)

    # Process specs
    specs = product_data.get('specs', [])

    if isinstance(specs, list):
        for group_item in specs:
            if "group" in group_item:
                # Group header row (merged across all 3 columns)
                row_cells = table.add_row().cells
                row_cells[0].merge(row_cells[2])
                p = row_cells[0].paragraphs[0]
                p.text = ""
                _add_bold_run(p, group_item["group"], size=Pt(10), font_name='Times New Roman')

            for spec in group_item.get("specs", []):
                _add_spec_row(table, spec.get("name", ""), spec.get("value", ""))

    elif isinstance(specs, dict):
        for key, value in specs.items():
            _add_spec_row(table, key, value)

    # Spacer after table
    document.add_paragraph('')


def _add_spec_row(table, name: str, value: str):
    """Add a single specification row with value/unit extraction."""
    row_cells = table.add_row().cells

    # Column 1: Name
    p0 = row_cells[0].paragraphs[0]
    p0.text = ""
    _add_run(p0, str(name), size=Pt(10), font_name='Times New Roman')

    # Extract value and unit
    display_val, unit = _extract_value_and_unit(str(name), str(value))

    # Column 2: Value (might have multiple lines)
    p1 = row_cells[1].paragraphs[0]
    p1.text = ""
    if "\n" in display_val:
        lines = display_val.split("\n")
        _add_run(p1, lines[0].strip(), size=Pt(10), font_name='Times New Roman')
        for line in lines[1:]:
            new_p = row_cells[1].add_paragraph()
            _add_run(new_p, line.strip(), size=Pt(10), font_name='Times New Roman')
    else:
        _add_run(p1, display_val, size=Pt(10), font_name='Times New Roman')

    # Column 3: Unit
    p2 = row_cells[2].paragraphs[0]
    p2.text = ""
    if unit:
        _add_run(p2, unit, size=Pt(10), font_name='Times New Roman')


if __name__ == "__main__":
    # Test with formal structure
    mock_products = [
        {
            "product_name": "Системный блок",
            "specs": [
                {
                    "group": "Корпус",
                    "specs": [
                        {"name": "тип", "value": "Мини-ПК"},
                        {"name": "Цвет", "value": "Черный"},
                        {"name": "Количество портов USB 2.0 Type A", "value": "не менее 2 шт"},
                        {"name": "Мощность блока питания", "value": "не менее 120 Вт"},
                    ]
                },
                {
                    "group": "Процессор",
                    "specs": [
                        {"name": "Тип Intel Core i7", "value": "не менее 1 шт"},
                        {"name": "Число ядер процессора", "value": "не менее 20 шт"},
                        {"name": "Частота процессора", "value": "не менее 2.1 ГГц"},
                        {"name": "Кэш-память L3", "value": "не менее 25 MB"},
                    ]
                }
            ]
        }
    ]
    mock_metadata = {
        "product_title": "системные блоки",
        "zakazchik": "ООО «Тест»",
        "quantity": 30,
        "quantity_text": "тридцать"
    }
    generate_tz_document('base_template.docx', 'test_tz.docx', mock_products, mock_metadata)
