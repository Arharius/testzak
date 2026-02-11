from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def create_base_template():
    doc = Document()

    # Page Layout (A4, Narrow Margins)
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # Define Styles
    styles = doc.styles

    style_normal = styles['Normal']
    font = style_normal.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # Heading 1
    style_h1 = styles['Heading 1']
    font = style_h1.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    font.bold = True
    style_h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_h1.paragraph_format.space_after = Pt(12)

    # Title
    doc.add_heading('ТЕХНИЧЕСКОЕ ЗАДАНИЕ', level=1)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('на поставку оборудования')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    doc.add_paragraph('')  # Spacing

    # Placeholder for dynamic content
    doc.add_paragraph('{{PRODUCT_TABLES_PLACEHOLDER}}')

    doc.save('base_template.docx')
    print("Template 'base_template.docx' created successfully.")


if __name__ == "__main__":
    create_base_template()
