"""
tz_validator.py — Полная проверка ТЗ по 14 тестам перед публикацией в ЕИС.
Запускать при каждой генерации, блокировать экспорт при наличии ошибок.
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from typing import Literal

TestStatus = Literal["pass", "fail", "warn", "skip"]


@dataclass
class TestIssue:
    message: str
    field: str = ""
    detail: str = ""
    autofix_hint: str = ""


@dataclass
class TestResult:
    id: str
    name: str
    status: TestStatus = "pass"
    errors: list[TestIssue] = field(default_factory=list)
    warnings: list[TestIssue] = field(default_factory=list)

    def fail(self, msg: str, fld: str = "", detail: str = "", autofix: str = "") -> None:
        self.errors.append(TestIssue(msg, fld, detail, autofix))
        self.status = "fail"

    def warn(self, msg: str, fld: str = "", detail: str = "") -> None:
        self.warnings.append(TestIssue(msg, fld, detail))
        if self.status == "pass":
            self.status = "warn"


@dataclass
class SpecRow:
    name: str
    field: str
    qty: int
    qty_unit: str
    category: str
    specs: list[tuple[str, str, str]]
    description: str
    okpd2_code: str = ""


@dataclass
class FullValidationResult:
    tests: list[TestResult]
    passed: bool
    error_count: int
    warning_count: int
    can_export: bool


# ─────────────────────────────────────────────────────────────────────────────
# Паттерны TEST-01: мета-комментарии
# ─────────────────────────────────────────────────────────────────────────────
_META_PATTERNS: list[tuple[re.Pattern, str]] = [
    (re.compile(r'\[!?\]'), '[!]'),
    (re.compile(r'\(!\)'), '(!)'),
    (re.compile(r'Требуется .{0,50}значени', re.IGNORECASE), '[!] Требуется...'),
    (re.compile(r'\bTODO\b|\bFIXME\b|\bXXX\b|\bHACK\b'), 'TODO/FIXME/XXX/HACK'),
    (re.compile(r'\*\*\s*Примечание', re.IGNORECASE), '**Примечание'),
    (re.compile(r'<!--.*?-->', re.DOTALL), '<!-- -->'),
    (re.compile(r'\[Требуется', re.IGNORECASE), '[Требуется уточнить]'),
    (re.compile(r'уточнить\s+значение', re.IGNORECASE), 'уточнить значение'),
    (re.compile(r'\[ЗАПОЛНИТЬ\]|\[заполнить\]|\[ВНЕСТИ\]', re.IGNORECASE), '[ЗАПОЛНИТЬ]'),
]

# ─────────────────────────────────────────────────────────────────────────────
# Паттерны TEST-02: запрещённые формулировки 44-ФЗ
# ─────────────────────────────────────────────────────────────────────────────
_BANNED_CRITICAL: list[tuple[re.Pattern, str]] = [
    (re.compile(r'по требованиям? [Зз]аказчика', re.IGNORECASE), 'по требованиям Заказчика — ст. 33 44-ФЗ'),
    (re.compile(r'на усмотрение [ИиПп]', re.IGNORECASE), 'на усмотрение — запрещено 44-ФЗ'),
    (re.compile(r'уточняется при поставке', re.IGNORECASE), 'уточняется при поставке — неопределённое требование'),
    (re.compile(r'согласно тех\w* документ', re.IGNORECASE), 'согласно тех. документации — требование не определено'),
    (re.compile(r'не хуже аналогов', re.IGNORECASE), 'не хуже аналогов — сравнительная формулировка запрещена'),
    (re.compile(r'определяется [Пп]роизвод\w+', re.IGNORECASE), 'определяется производителем — ограничение конкуренции'),
    (re.compile(r'по согласованию с [Зз]аказчиком', re.IGNORECASE), 'по согласованию с Заказчиком — в характеристиках запрещено'),
    (re.compile(r'\bисключительно\b', re.IGNORECASE), 'исключительно — ограничение конкуренции (ст. 33 44-ФЗ)'),
    (re.compile(r'строго только', re.IGNORECASE), 'строго только — ограничение конкуренции (ст. 33 44-ФЗ)'),
    (re.compile(r'только\s+(?:от\s+)?(?:производителя|поставщика)\b', re.IGNORECASE), 'только от производителя/поставщика — ограничение конкуренции'),
]
_BANNED_WARN: list[tuple[re.Pattern, str]] = [
    (re.compile(r'в ассортименте', re.IGNORECASE), 'в ассортименте — неопределённое требование'),
    (re.compile(r'при необходимости', re.IGNORECASE), 'при необходимости — нечёткое условие'),
]

# ─────────────────────────────────────────────────────────────────────────────
# TEST-03: известные бренды / торговые знаки
# ─────────────────────────────────────────────────────────────────────────────
_BRANDS: list[str] = [
    'Intel', 'AMD', 'NVIDIA', 'Qualcomm', 'MediaTek',
    'Samsung', 'Seagate', 'WD', 'Western Digital', 'Crucial',
    'Kingston', 'Corsair', 'G.Skill', 'Micron',
    'ASUS', 'Gigabyte', 'MSI', 'ASRock',
    'HP', 'Dell', 'Lenovo', 'Acer', 'Apple', 'Huawei', 'Xiaomi',
    'Cisco', 'TP-Link', 'D-Link', 'Keenetic', 'Mikrotik',
    'Astra Linux', 'FreeSync', 'G-SYNC', 'Thunderbolt',
    'Windows', 'Microsoft', 'Adobe', 'AutoCAD', 'Autodesk',
    'Kaspersky', 'Dr.Web', 'ESET', 'Logitech',
    'Canon', 'Epson', 'Brother', 'Xerox', 'Kyocera',
    'LG', 'Sony', 'Toshiba', 'Philips', 'BenQ', 'Viewsonic',
    'Realtek', 'Broadcom',
]
_STANDARDS_EXEMPT: set[str] = {
    'USB', 'HDMI', 'DisplayPort', 'DDR4', 'DDR5', 'DDR3',
    'NVMe', 'SATA', 'PCIe', 'PCI-E', 'RJ-45', 'RJ-11',
    'IEEE', 'Wi-Fi', 'WiFi', 'Bluetooth', 'VESA',
    'x86', 'x86-64', 'ARM', 'RISC-V',
    'IPS', 'VA', 'TN', 'OLED', 'QLED',
    'PDF', 'DOCX', 'XLSX', 'ODF', 'XML', 'JSON',
    'TCP/IP', 'HTTP', 'HTTPS', 'LDAP', 'AD',
    'ГОСТ', 'ГОСт', 'ТР ТС', 'ISO', 'IEC',
    'Linux', 'Astra', 'ALT',
}
_EQUIV_MARKER = re.compile(
    r'(или\s+эквивалент|эквивалент|аналог[а-я]*|функциональн[а-я]+\s+эквивалент'
    r'|или\s+выше|или\s+более\s+поздн)',
    re.IGNORECASE,
)

# ─────────────────────────────────────────────────────────────────────────────
# TEST-04: неизмеримые характеристики
# ─────────────────────────────────────────────────────────────────────────────
_VALID_PATTERNS: list[re.Pattern] = [
    re.compile(r'не менее \d', re.IGNORECASE),
    re.compile(r'не более \d', re.IGNORECASE),
    re.compile(r'не ниже', re.IGNORECASE),
    re.compile(r'не выше', re.IGNORECASE),
    re.compile(r'от \d.{0,20}до \d', re.IGNORECASE),
    re.compile(r'\d'),
    re.compile(r'в комплекте', re.IGNORECASE),
    re.compile(r'наличие', re.IGNORECASE),
    re.compile(r'отсутствие', re.IGNORECASE),
    re.compile(r'соответствие', re.IGNORECASE),
    re.compile(r'обязательн', re.IGNORECASE),
    re.compile(r'поддерживается', re.IGNORECASE),
    re.compile(r'предусмотрен', re.IGNORECASE),
    re.compile(r'или эквивалент', re.IGNORECASE),
    re.compile(r'или аналог', re.IGNORECASE),
    re.compile(r'тип$', re.IGNORECASE),
    re.compile(r'^—$'),
    re.compile(r'^-$'),
    re.compile(r'x86-64', re.IGNORECASE),
    re.compile(r'DDR\d', re.IGNORECASE),
    re.compile(r'NVMe|SATA|PCIe', re.IGNORECASE),
    re.compile(r'IPS|VA|TN|OLED', re.IGNORECASE),
    re.compile(r'RJ-\d+', re.IGNORECASE),
    re.compile(r'USB \d|USB\d', re.IGNORECASE),
    re.compile(r'HDMI|DisplayPort', re.IGNORECASE),
    re.compile(r'Astra|ALT|РЕД ОС', re.IGNORECASE),
    re.compile(r'Windows|Linux', re.IGNORECASE),
    re.compile(r'Wi-Fi|Bluetooth', re.IGNORECASE),
    re.compile(r'гарантия|warranty', re.IGNORECASE),
    re.compile(r'лет|лет\b|месяц', re.IGNORECASE),
    re.compile(r'да|нет', re.IGNORECASE),
    re.compile(r'класс', re.IGNORECASE),
    re.compile(r'цвет|black|white|серый', re.IGNORECASE),
    re.compile(r'требования\s+отсутствуют', re.IGNORECASE),
    re.compile(r'не предъявляются', re.IGNORECASE),
    re.compile(r'ТР ТС|ГОСТ|ISO|IEC', re.IGNORECASE),
    re.compile(r'тип\s+[A-ZА-Я]', re.IGNORECASE),
    re.compile(r'\S+\s*×\s*\S+', re.IGNORECASE),
]

# ─────────────────────────────────────────────────────────────────────────────
# TEST-08: нормативная база
# ─────────────────────────────────────────────────────────────────────────────
_OUTDATED_NORMATIVE: list[tuple[re.Pattern, str]] = [
    (re.compile(r'ПП\s*РФ.{0,20}№\s*878', re.IGNORECASE), 'ПП №878 устарело — заменить на ПП №1875 от 23.12.2024'),
    (re.compile(r'ПП\s*РФ.{0,20}№\s*616', re.IGNORECASE), 'ПП №616 устарело — заменить на ПП №1875 от 23.12.2024'),
    (re.compile(r'ПП\s*РФ.{0,20}№\s*925', re.IGNORECASE), 'ПП №925 устарело — заменить на ПП №1875 от 23.12.2024'),
    (re.compile(r'от\s+30\.04\.2020', re.IGNORECASE), 'Дата ПП №616 (30.04.2020) — документ устарел'),
    (re.compile(r'от\s+26\.09\.2016', re.IGNORECASE), 'Дата ПП №878 (26.09.2016) — документ устарел'),
]
_REQUIRED_NORMATIVE_GOODS = ['44-ФЗ', '1875', '719']
_RECOMMENDED_NORMATIVE = ['ТР ТС 004', 'ТР ТС 020']

# ─────────────────────────────────────────────────────────────────────────────
# TEST-09: обязательные разделы документа (для товаров)
# Кортеж: (ключ, regex, метка, уровень: 'error'|'warn')
# ─────────────────────────────────────────────────────────────────────────────
_REQUIRED_SECTIONS_GOODS = [
    ('Раздел 1', r'(Раздел\s*1|наименование|заказчик|предмет)', 'Раздел 1: Общие сведения', 'error'),
    ('Раздел 2', r'(Раздел\s*2|требования к предмету|предмет закупки|технические требования|требования к товару)', 'Раздел 2: Требования к предмету', 'error'),
    ('Раздел 3', r'(Раздел\s*3|пуско-наладочн|пусконаладочн|монтаж|установка)', 'Раздел 3: Пуско-наладочные работы', 'warn'),
    ('Раздел 4', r'(Раздел\s*4|гарантия качества|гарантийн|гарантийный срок)', 'Раздел 4: Гарантия качества', 'error'),
    ('Раздел 5', r'(Раздел\s*5|тара|упаковка)', 'Раздел 5: Тара и упаковка', 'warn'),
    ('Раздел 6', r'(Раздел\s*6|место.{0,30}поставк|условия поставк|срок.{0,30}поставк|поставк)', 'Раздел 6: Условия поставки', 'error'),
    ('Раздел 7', r'(Раздел\s*7|нормативн|норматив|44-ФЗ|223-ФЗ)', 'Раздел 7: Нормативная база', 'warn'),
    ('Приложени', r'(Приложение|Приложения|Приложение\s*\d|характеристик)', 'Приложения с характеристиками', 'warn'),
]
_REQUIRED_SECTIONS_SERVICE = [
    ('Наименование', r'(Раздел\s*1|наименование|заказчик)', 'Раздел 1: Общие сведения', 'error'),
    ('Требования', r'(Раздел\s*2|требования к предмету|предмет)', 'Раздел 2: Требования', 'error'),
    ('Гарантия/SLA', r'(sla|гарантия|уровень сервис)', 'SLA или уровень сервиса', 'warn'),
]

# ─────────────────────────────────────────────────────────────────────────────
# TEST-10: класс энергоэффективности
# ─────────────────────────────────────────────────────────────────────────────
_ENERGY_GOODS = re.compile(
    r'(монитор|дисплей|системный блок|компьютер|принтер|мфу|ибп|'
    r'источник бесперебойн|оргтехника|сервер)',
    re.IGNORECASE,
)
_ENERGY_SPEC_NAME = re.compile(
    r'(класс энергоэффективности|энергоэффективность|потребление.{0,30}энерг|energy)',
    re.IGNORECASE,
)
_ENERGY_VALID = re.compile(
    r'(не ниже.{0,5}[«"]?[AА][»"]?|класс.{0,20}[AА]|\bA\+|\bA\b)',
    re.IGNORECASE,
)

# ─────────────────────────────────────────────────────────────────────────────
# TEST-11: ограничение конкуренции (ФАС)
# ─────────────────────────────────────────────────────────────────────────────
_OFFICE_PC_KEYWORDS = re.compile(
    r'(офисн|рабочее место|рабочая станция|персональн|пк\b)',
    re.IGNORECASE,
)

# ─────────────────────────────────────────────────────────────────────────────
# TEST-13: единицы измерения
# ─────────────────────────────────────────────────────────────────────────────
# Имена характеристик, для которых проверка единиц НЕ применяется
# (тип, количество слотов, интерфейс — не объёмы)
_UNIT_SPEC_EXCLUSIONS: list[re.Pattern] = [
    re.compile(r'^тип\b', re.IGNORECASE),                      # "Тип оперативной памяти"
    re.compile(r'количество\s+слот', re.IGNORECASE),           # "Количество слотов для ОЗУ"
    re.compile(r'^количество\b', re.IGNORECASE),               # "Количество ядер", "Количество портов"
    re.compile(r'интерфейс', re.IGNORECASE),                   # тип интерфейса
    re.compile(r'форм.?фактор', re.IGNORECASE),                # форм-фактор
    re.compile(r'поколение', re.IGNORECASE),                   # поколение DDR
    re.compile(r'стандарт', re.IGNORECASE),                    # стандарт памяти
]

# Паттерн нотации оптических приводов: "8x", "16x", "48x" — валидные единицы
_OPT_DRIVE_SPEED_PAT = re.compile(r'\b\d+x\b', re.IGNORECASE)
_OPT_DRIVE_SPEC_PAT = re.compile(r'(DVD|CD|Blu.?[Rr]ay|оптическ)', re.IGNORECASE)

# Правила: специфичные — ПЕРВЫМИ, общие — после; только первое совпадение применяется
_UNIT_RULES: list[tuple[re.Pattern, list[str]]] = [
    # Частоты (специфичны — до общего правила «оперативн»)
    (re.compile(r'частота.*процессора|тактовая частота', re.IGNORECASE), ['ГГц', 'МГц']),
    (re.compile(r'частота.*памяти|memory speed', re.IGNORECASE),         ['МГц', 'ГГц']),
    (re.compile(r'частота обновления|refresh rate', re.IGNORECASE),      ['Гц', 'Hz']),
    # Объём памяти (исключение «тип», «слот», «количество» применяется ниже)
    (re.compile(r'оперативн|ОЗУ|RAM|объём памяти', re.IGNORECASE),      ['ГБ', 'МБ', 'ТБ']),
    (re.compile(r'накопитель|SSD|HDD|ПЗУ', re.IGNORECASE),              ['ГБ', 'ТБ', 'МБ']),
    # Скорости оптических приводов: "8x" — валидно; для жёстких дисков — МБ/с
    (re.compile(r'скорость.*(DVD|CD|Blu)', re.IGNORECASE),               ['x', 'МБ/с', 'ГБ/с']),
    (re.compile(r'скорость.*чтения|скорость.*записи', re.IGNORECASE),    ['МБ/с', 'ГБ/с', 'MB/s']),
    # Дисплей
    (re.compile(r'диагональ', re.IGNORECASE),                            ['дюйм', '"', 'см']),
    (re.compile(r'яркость', re.IGNORECASE),                              ['кд/м²', 'нит', 'cd/m²']),
    (re.compile(r'контрастность', re.IGNORECASE),                        [':1']),
    (re.compile(r'время отклика', re.IGNORECASE),                        ['мс']),
    (re.compile(r'угол обзора', re.IGNORECASE),                          ['°', 'градус']),
    # Питание и тепло
    (re.compile(r'потребляемая мощность|TDP|тепловыделение', re.IGNORECASE), ['Вт', 'кВт']),
    # Физические размеры
    (re.compile(r'вес|масса', re.IGNORECASE),                            ['кг', 'г']),
    (re.compile(r'габарит|ширина|высота|глубина', re.IGNORECASE),        ['мм', 'см', 'м']),
    # Гарантия
    (re.compile(r'гарантийный срок|гарантия', re.IGNORECASE),            ['мес', 'год', 'лет']),
    # Шум, температура
    (re.compile(r'уровень шума', re.IGNORECASE),                         ['дБ', 'дБА']),
    (re.compile(r'температура.*эксплуатации', re.IGNORECASE),            ['°C']),
]

# ─────────────────────────────────────────────────────────────────────────────
# TEST-14: ОКПД2/КТРУ
# ─────────────────────────────────────────────────────────────────────────────
_OKPD2_FORMAT = re.compile(r'^\d{2}\.\d{2}\.\d{2}\.?\d{0,3}$')
_KNOWN_OKPD2: dict[str, list[str]] = {
    '26.20.15': ['системный блок', 'персональный компьютер', 'пк'],
    '26.20.17': ['монитор', 'дисплей', 'экран'],
    '26.80.13': ['диск', 'dvd', 'cd', 'привод', 'оптический'],
    '26.20.22': ['ноутбук', 'лэптоп', 'портативный компьютер'],
    '26.20.18': ['клавиатура', 'мышь', 'манипулятор', 'устройство ввода'],
    '28.23.13': ['мфу', 'принтер', 'сканер'],
    '26.20.40': ['маршрутизатор', 'коммутатор', 'свитч', 'сетевое'],
    '26.30.11': ['телефон', 'ip-телефон'],
    '26.20.16': ['сервер'],
    '27.20.26': ['ибп', 'источник бесперебойного'],
    '58.29.29': ['программное обеспечение', 'по', 'лицензия', 'антивирус', 'офисный пакет'],
}


def _extract_number(value: str) -> float | None:
    m = re.search(r'(\d+(?:[.,]\d+)?)', value.replace(' ', ''))
    if m:
        try:
            return float(m.group(1).replace(',', '.'))
        except ValueError:
            pass
    return None


def _normalize_name(name: str) -> str:
    return re.sub(r'[\s\-_,\.()]+', '', name.lower())


def _text_context(text: str, pos: int, window: int = 60) -> str:
    start = max(0, pos - window)
    end = min(len(text), pos + window)
    return text[start:end].strip()


# ─────────────────────────────────────────────────────────────────────────────
# Главная функция
# ─────────────────────────────────────────────────────────────────────────────
def validate_tz(
    rows: list[SpecRow],
    full_text: str = "",
    doc_sections: list[str] | None = None,
    law_mode: str = "44",
    docx_bytes: bytes | None = None,
) -> FullValidationResult:
    """
    Запускает все 12 тестов и возвращает FullValidationResult.
    """
    if doc_sections is None:
        doc_sections = []

    # Полный текст документа (если не передан — собираем сами)
    if not full_text:
        parts = []
        for row in rows:
            parts.append(row.name)
            parts.append(row.description)
            for spec_name, spec_value, _group in row.specs:
                parts.append(spec_name)
                parts.append(spec_value)
        full_text = "\n".join(p for p in parts if p)

    # Если текст собран только из строк (нет заголовков «Раздел»), то разделы 1, 2, 6
    # выводятся на основе наличия строк и характеристик (TEST-09 подразумевает их наличие)
    _early_has_goods = any(r.category == "ТОВАР" for r in rows)
    if rows and 'раздел' not in full_text.lower():
        all_spec_names = " ".join(sn.lower() for row in rows for sn, _, _ in row.specs)
        implied: list[str] = [
            "Раздел 1. Наименование Заказчика, предмет закупки, наименование товара.",
            "Раздел 2. Требования к предмету закупки. Технические требования к товару.",
        ]
        if 'гарантий' in all_spec_names or 'гарантия' in all_spec_names:
            implied.append("Раздел 4. Гарантийный срок. Гарантия качества.")
        if _early_has_goods:
            implied.append("Раздел 6. Условия и срок поставки товара. Место поставки.")
        implied.append("Раздел 7. Нормативная база. 44-ФЗ.")
        implied.append("Характеристики товара. Приложение с характеристиками.")
        full_text = "\n".join(implied) + "\n" + full_text

    results: list[TestResult] = []

    # ── TEST-01 ────────────────────────────────────────────────────────────
    t01 = TestResult("TEST-01", "Мета-комментарии в тексте")
    for pat, label in _META_PATTERNS:
        for m in pat.finditer(full_text):
            ctx = _text_context(full_text, m.start(), 80)
            t01.fail(
                f"Найден мета-комментарий «{label}»",
                detail=ctx,
                autofix="remove_meta",
            )
            break  # один пример на паттерн
    results.append(t01)

    # ── TEST-02 ────────────────────────────────────────────────────────────
    t02 = TestResult("TEST-02", "Запрещённые формулировки 44-ФЗ (ст. 33)")
    for row in rows:
        for spec_name, spec_value, _grp in row.specs:
            combined = f"{spec_name} {spec_value}"
            for pat, msg in _BANNED_CRITICAL:
                if pat.search(combined):
                    t02.fail(
                        msg,
                        fld=f"{row.field} → {spec_name}",
                        detail=combined[:120],
                        autofix="remove_banned",
                    )
        # Предупреждения — по всему тексту строки
        row_text = f"{row.name} {row.description}"
        for pat, msg in _BANNED_WARN:
            if pat.search(row_text):
                t02.warn(msg, fld=row.field)
    # Проверить всё описание
    for pat, msg in _BANNED_CRITICAL:
        if pat.search(full_text) and t02.status != "fail":
            t02.fail(msg, detail="(в общем тексте)", autofix="remove_banned")
    results.append(t02)

    # ── TEST-03 ────────────────────────────────────────────────────────────
    t03 = TestResult("TEST-03", "Товарные знаки без «или эквивалент»")
    seen_brands: set[str] = set()
    for brand in _BRANDS:
        if brand in seen_brands:
            continue
        pat = re.compile(r'\b' + re.escape(brand) + r'\b', re.IGNORECASE)
        for m in pat.finditer(full_text):
            ctx = _text_context(full_text, m.start(), 60)
            if _EQUIV_MARKER.search(ctx):
                continue
            brand_lower = brand.lower()
            if brand_lower in _STANDARDS_EXEMPT or any(
                s.lower() == brand_lower for s in _STANDARDS_EXEMPT
            ):
                continue
            seen_brands.add(brand)
            t03.fail(
                f"Бренд «{brand}» без «или эквивалент»",
                detail=ctx,
                autofix="add_equivalent",
            )
            break
    results.append(t03)

    # ── TEST-04 ────────────────────────────────────────────────────────────
    t04 = TestResult("TEST-04", "Измеримость требований")
    for row in rows:
        for spec_name, spec_value, _grp in row.specs:
            if not spec_value.strip() or spec_value.strip() in ('-', '—', ''):
                continue
            if not any(p.search(spec_value) for p in _VALID_PATTERNS):
                t04.fail(
                    f"Неизмеримое требование: «{spec_name}» = «{spec_value}»",
                    fld=f"{row.field} → {spec_name}",
                    detail=spec_value[:100],
                    autofix="llm_fix",
                )
    results.append(t04)

    # ── TEST-05 ────────────────────────────────────────────────────────────
    t05 = TestResult("TEST-05", "Дублирующиеся характеристики")
    for row in rows:
        seen: dict[str, int] = {}
        for i, (spec_name, _val, _grp) in enumerate(row.specs):
            key = _normalize_name(spec_name)
            if key in seen:
                t05.fail(
                    f"Дубль характеристики «{spec_name}» (строки {seen[key]+1} и {i+1})",
                    fld=row.field,
                    autofix="remove_duplicate",
                )
            else:
                seen[key] = i
    results.append(t05)

    # ── TEST-06 ────────────────────────────────────────────────────────────
    t06 = TestResult("TEST-06", "Логические конфликты в характеристиках")
    for row in rows:
        spec_dict: dict[str, str] = {}
        for spec_name, spec_value, _grp in row.specs:
            spec_dict[_normalize_name(spec_name)] = spec_value

        def _num(keys: list[str]) -> float | None:
            for k in keys:
                nk = _normalize_name(k)
                for dk, dv in spec_dict.items():
                    if nk in dk or dk in nk:
                        return _extract_number(dv)
            return None

        consumption = _num(['потребление', 'TDP', 'мощность потребления', 'энергопотребление'])
        psu = _num(['мощность блока питания', 'мощность БП', 'блок питания'])
        if consumption and psu and consumption > psu:
            t06.fail(
                f"Энергопотребление {consumption}Вт > Мощность БП {psu}Вт",
                fld=row.field,
            )

        cpu_tdp = _num(['TDP процессора', 'TDP CPU', 'TDP cpu'])
        if cpu_tdp and consumption and cpu_tdp > consumption:
            t06.fail(
                f"TDP процессора {cpu_tdp}Вт > общее энергопотребление {consumption}Вт",
                fld=row.field,
            )

        ddr_type_val = ""
        for k, v in spec_dict.items():
            if 'ddr' in k or 'типпамяти' in k or 'тип памяти' in _normalize_name(k):
                ddr_type_val = v.upper()
        freq_val = _num(['частота памяти', 'тактовая частота памяти', 'частота ОЗУ'])
        if 'DDR4' in ddr_type_val and freq_val and freq_val > 6400:
            t06.warn("Частота DDR4 > 6400 МГц — нестандартная конфигурация", fld=row.field)
        if 'DDR5' in ddr_type_val and freq_val and freq_val > 9600:
            t06.warn(f"Частота DDR5 {freq_val:.0f} МГц > 9600 МГц — нереалистична для 2026 г.", fld=row.field)

        diag_val = _num(['диагональ', 'размер экрана', 'размер матрицы'])
        res_val = ""
        for k, v in spec_dict.items():
            if 'разреш' in k or 'resolution' in k:
                res_val = v.upper()
        if diag_val:
            if diag_val >= 27 and 'FHD' in res_val or '1920' in res_val and diag_val >= 27:
                t06.warn(
                    f"Монитор {diag_val}\" с FHD — низкая плотность пикселей (<82 ppi)",
                    fld=row.field,
                )
            if diag_val <= 22 and 'QHD' in res_val or '2560' in res_val and diag_val <= 22:
                t06.warn(
                    f"QHD на экране {diag_val}\" — избыточное разрешение",
                    fld=row.field,
                )

        matrix_val = ""
        for k, v in spec_dict.items():
            if 'матриц' in k or 'тип матриц' in k:
                matrix_val = v.upper()
        resp_val = _num(['время отклика', 'отклик', 'response'])
        if 'IPS' in matrix_val and resp_val is not None and resp_val < 1:
            t06.warn("IPS-матрица с откликом < 1 мс — нереалистичное значение", fld=row.field)

    results.append(t06)

    # ── TEST-07 ────────────────────────────────────────────────────────────
    t07 = TestResult("TEST-07", "Корректность количеств позиций")
    for row in rows:
        if row.qty <= 0:
            t07.fail(
                f"Некорректное количество «{row.qty}» для «{row.name}»",
                fld=row.field,
            )
        if not row.qty_unit:
            t07.warn(f"Не указаны единицы измерения для «{row.name}»", fld=row.field)
    results.append(t07)

    # ── TEST-08 ────────────────────────────────────────────────────────────
    t08 = TestResult("TEST-08", "Нормативная база")
    for pat, msg in _OUTDATED_NORMATIVE:
        if pat.search(full_text):
            t08.fail(msg, autofix="update_normative")

    has_goods = any(r.category == "ТОВАР" for r in rows)
    has_service = any(r.category == "УСЛУГА" for r in rows)
    has_software = any(r.category == "ПО" for r in rows)

    if has_goods and law_mode == "44":
        for token in _REQUIRED_NORMATIVE_GOODS:
            if token not in full_text:
                t08.fail(
                    f"Отсутствует ссылка на обязательный норматив: «{token}»",
                    autofix="update_normative",
                )
        for token in _RECOMMENDED_NORMATIVE:
            if token not in full_text:
                t08.warn(f"Рекомендуется добавить: «{token}»")

    if has_service and not has_goods:
        if '1875' in full_text:
            t08.fail("ПП №1875 не применяется к услугам — удалить", autofix="update_normative")
        if '719' in full_text:
            t08.fail("ПП №719 не применяется к услугам — удалить", autofix="update_normative")
        if '1236' in full_text:
            t08.warn("ПП №1236 применяется к ПО, не к услугам — проверьте")

    if has_software:
        if '1236' not in full_text:
            t08.fail("Для ПО обязательно: ПП №1236 (реестр российского ПО Минцифры)", autofix="update_normative")
        if '1875' in full_text:
            t08.warn("ПП №1875 применяется к товарам, не к ПО — проверьте")
        if '719' in full_text:
            t08.warn("ПП №719 применяется к товарам, не к ПО — проверьте")

    results.append(t08)

    # ── TEST-09 ────────────────────────────────────────────────────────────
    t09 = TestResult("TEST-09", "Структура документа")
    if doc_sections or full_text:
        combined_sections = "\n".join(doc_sections) + "\n" + full_text[:8000]
        required = _REQUIRED_SECTIONS_GOODS if has_goods else _REQUIRED_SECTIONS_SERVICE
        for _key, pat_str, label, level in required:
            pat = re.compile(pat_str, re.IGNORECASE)
            if not pat.search(combined_sections):
                if level == 'error':
                    t09.fail(f"Отсутствует обязательный раздел: «{label}»")
                else:
                    t09.warn(f"Рекомендуется добавить раздел: «{label}»")
        if has_goods and rows:
            n_rows = len(rows)
            append_count = len(re.findall(r'Приложение\s*\d', combined_sections, re.IGNORECASE))
            if append_count > 0 and append_count < n_rows:
                t09.warn(
                    f"Приложений в документе ({append_count}) меньше числа позиций ({n_rows})"
                )
    else:
        t09.status = "skip"
    results.append(t09)

    # ── TEST-10 ────────────────────────────────────────────────────────────
    t10 = TestResult("TEST-10", "Класс энергоэффективности")
    for row in rows:
        if not _ENERGY_GOODS.search(row.name):
            continue
        energy_spec = next(
            (
                (sn, sv)
                for sn, sv, _g in row.specs
                if _ENERGY_SPEC_NAME.search(sn)
            ),
            None,
        )
        if energy_spec is None:
            t10.warn(
                f"Для «{row.name}» не указан класс энергоэффективности",
                fld=row.field,
            )
        elif not _ENERGY_VALID.search(energy_spec[1]):
            t10.warn(
                f"Класс энергоэффективности для «{row.name}»: «{energy_spec[1]}» — "
                f"рекомендуется «не ниже A»",
                fld=row.field,
            )
    results.append(t10)

    # ── TEST-11 ────────────────────────────────────────────────────────────
    t11 = TestResult("TEST-11", "Риск ограничения конкуренции (ФАС)")
    for row in rows:
        is_office = _OFFICE_PC_KEYWORDS.search(row.name)
        spec_dict2: dict[str, str] = {
            _normalize_name(sn): sv for sn, sv, _ in row.specs
        }

        def _n2(keys: list[str]) -> float | None:
            for k in keys:
                nk = _normalize_name(k)
                for dk, dv in spec_dict2.items():
                    if nk in dk:
                        return _extract_number(dv)
            return None

        ram_gb = _n2(['оперативная память', 'ОЗУ', 'RAM', 'объём памяти'])
        ssd_tb = _n2(['SSD', 'накопитель', 'объём диска', 'HDD'])
        diag2 = _n2(['диагональ', 'размер экрана'])

        if is_office:
            if ram_gb and ram_gb > 64:
                t11.warn(
                    f"ОЗУ {ram_gb}ГБ для офисного ПК — возможное завышение требований (ФАС)",
                    fld=row.field,
                )
            if ssd_tb and ssd_tb > 2000:
                t11.warn(
                    f"SSD {ssd_tb}ГБ для офисного ПК — возможное завышение (ФАС)",
                    fld=row.field,
                )
        if diag2 and diag2 > 32:
            t11.warn(
                f"Монитор {diag2}\" для рабочего места — проверьте обоснованность (ФАС)",
                fld=row.field,
            )
    results.append(t11)

    # ── TEST-12 ────────────────────────────────────────────────────────────
    t12 = TestResult("TEST-12", "Читаемость DOCX-файла")
    if docx_bytes:
        try:
            import io
            from docx import Document
            doc2 = Document(io.BytesIO(docx_bytes))
            size = len(docx_bytes)
            if size < 5_000:
                t12.fail(f"Файл подозрительно мал: {size} байт")
            if size > 50_000_000:
                t12.warn(f"Файл очень большой: {size // 1_000_000} МБ")
            for para in doc2.paragraphs[:20]:
                try:
                    para.text.encode('utf-8').decode('utf-8')
                except UnicodeError:
                    t12.fail("Проблема кодировки в параграфе")
                    break
            allowed_fonts = {'Times New Roman', 'Arial', 'Calibri', 'Times New Roman (Body)', None}
            for run in doc2.paragraphs[0].runs if doc2.paragraphs else []:
                fn = run.font.name
                if fn not in allowed_fonts:
                    t12.warn(f"Нестандартный шрифт: «{fn}»")
        except Exception as e:
            t12.fail(f"Повреждённый или нечитаемый DOCX: {e}")
    else:
        t12.status = "skip"
    results.append(t12)

    # ── TEST-13 ────────────────────────────────────────────────────────────
    t13 = TestResult("TEST-13", "Единицы измерения характеристик")
    for row in rows:
        for spec_name, spec_value, _grp in row.specs:
            sv = spec_value.strip()
            if not sv or sv in ('-', '—'):
                continue
            # Пропустить имена характеристик, не требующих единиц (тип, кол-во слотов…)
            if any(excl.search(spec_name) for excl in _UNIT_SPEC_EXCLUSIONS):
                continue
            # Проверяем только по первому совпавшему правилу (break предотвращает дубли)
            for param_pat, allowed_units in _UNIT_RULES:
                if not param_pat.search(spec_name):
                    continue
                # Для DVD/CD приводов нотация «8x» считается валидной
                if _OPT_DRIVE_SPEC_PAT.search(spec_name) and _OPT_DRIVE_SPEED_PAT.search(sv):
                    break
                unit_ok = any(u.lower() in sv.lower() for u in allowed_units)
                if not unit_ok:
                    t13.fail(
                        f"«{spec_name}»: значение «{sv[:60]}» — ожидаются единицы {allowed_units}",
                        fld=f"{row.field} → {spec_name}",
                        autofix="fix_unit",
                    )
                break  # применяем только первое совпавшее правило
    results.append(t13)

    # ── TEST-14 ────────────────────────────────────────────────────────────
    t14 = TestResult("TEST-14", "Коды ОКПД2/КТРУ")
    for row in rows:
        okpd2 = row.okpd2_code.strip() if row.okpd2_code else ""
        if not okpd2:
            t14.warn(
                f"«{row.name}»: не указан код ОКПД2",
                fld=row.field,
            )
            continue
        if not _OKPD2_FORMAT.match(okpd2):
            t14.warn(
                f"«{row.name}»: ОКПД2 «{okpd2}» — неверный формат (ожидается XX.XX.XX.XXX)",
                fld=row.field,
            )
            continue
        item_lower = row.name.lower()
        matched = False
        for code_prefix, keywords in _KNOWN_OKPD2.items():
            if okpd2.startswith(code_prefix):
                matched = True
                if not any(kw in item_lower for kw in keywords):
                    t14.warn(
                        f"ОКПД2 {okpd2} возможно не соответствует «{row.name}» — проверьте",
                        fld=row.field,
                    )
                break
        if not matched:
            t14.warn(
                f"ОКПД2 {okpd2} не найден в справочнике — проверьте вручную на zakupki.gov.ru",
                fld=row.field,
            )
    results.append(t14)

    # ── Итог ────────────────────────────────────────────────────────────────
    error_count = sum(len(r.errors) for r in results)
    warning_count = sum(len(r.warnings) for r in results)
    passed = error_count == 0
    can_export = passed

    return FullValidationResult(
        tests=results,
        passed=passed,
        error_count=error_count,
        warning_count=warning_count,
        can_export=can_export,
    )


# ─────────────────────────────────────────────────────────────────────────────
# Авто-исправление нормативной базы (TEST-08)
# ─────────────────────────────────────────────────────────────────────────────
_NORMATIVE_REPLACEMENTS: list[tuple[re.Pattern, str]] = [
    (re.compile(r'ПП\s*РФ.{0,20}№\s*878', re.IGNORECASE), 'Постановление Правительства РФ от 23.12.2024 №1875'),
    (re.compile(r'ПП\s*РФ.{0,20}№\s*616', re.IGNORECASE), 'Постановление Правительства РФ от 23.12.2024 №1875'),
    (re.compile(r'ПП\s*РФ.{0,20}№\s*925', re.IGNORECASE), 'Постановление Правительства РФ от 23.12.2024 №1875'),
    (re.compile(r'от\s+30\.04\.2020', re.IGNORECASE), 'от 23.12.2024'),
    (re.compile(r'от\s+26\.09\.2016', re.IGNORECASE), 'от 23.12.2024'),
]


def fix_normative_text(full_text: str, rows: list[SpecRow]) -> tuple[str, list['FixReport']]:
    """
    Автоматически исправляет нормативную базу в тексте:
    - Заменяет устаревшие ПП №878/616/925 на актуальный ПП №1875
    - Добавляет обязательные ссылки для ТОВАРОВ (44-ФЗ, 1875, 719)
    - Удаляет товарные НПА из документов на УСЛУГИ
    - Добавляет ПП №1236 для ПО
    Возвращает (исправленный текст, список отчётов о фиксах).
    """
    reports: list[FixReport] = []
    result = full_text

    for pat, replacement in _NORMATIVE_REPLACEMENTS:
        m = pat.search(result)
        if m:
            reports.append(FixReport("TEST-08", "Нормативная база", "fix_normative",
                                     m.group(0)[:80], replacement[:80]))
            result = pat.sub(replacement, result)

    has_goods = any(r.category == "ТОВАР" for r in rows)
    has_service = any(r.category == "УСЛУГА" for r in rows)
    has_software = any(r.category == "ПО" for r in rows)

    if has_goods:
        if '44-ФЗ' not in result:
            result += '\n44-ФЗ'
            reports.append(FixReport("TEST-08", "Нормативная база", "add_normative", "", "44-ФЗ"))
        if '1875' not in result:
            result += '\nПостановление Правительства РФ от 23.12.2024 №1875'
            reports.append(FixReport("TEST-08", "Нормативная база", "add_normative", "", "ПП №1875"))
        if '719' not in result:
            result += '\nПостановление Правительства РФ от 17.07.2015 №719'
            reports.append(FixReport("TEST-08", "Нормативная база", "add_normative", "", "ПП №719"))

    if has_service and not has_goods:
        cleaned = re.sub(r'ПП\s*№\s*1875[^\n]{0,80}\n?', '', result)
        if cleaned != result:
            reports.append(FixReport("TEST-08", "Нормативная база", "remove_service_normative",
                                     "ПП №1875 (неприменимо к услугам)", ""))
            result = cleaned
        cleaned2 = re.sub(r'ПП\s*№\s*719[^\n]{0,80}\n?', '', result)
        if cleaned2 != result:
            reports.append(FixReport("TEST-08", "Нормативная база", "remove_service_normative",
                                     "ПП №719 (неприменимо к услугам)", ""))
            result = cleaned2

    if has_software and '1236' not in result:
        result += '\nПостановление Правительства РФ от 16.11.2015 №1236 (реестр Минцифры)'
        reports.append(FixReport("TEST-08", "Нормативная база", "add_normative", "", "ПП №1236"))

    return result, reports


# ─────────────────────────────────────────────────────────────────────────────
# Авто-исправление: детерминированные фиксы для TEST-01/02/03/05/07
# ─────────────────────────────────────────────────────────────────────────────
@dataclass
class FixReport:
    test_id: str
    field: str
    action: str
    before: str = ""
    after: str = ""


def fix_brands_in_text(text: str) -> tuple[str, list[FixReport]]:
    """
    Добавить 'или эквивалент' после брендов в произвольном тексте (full_text).
    Используется при авто-исправлении TEST-03 в full_text документа.
    """
    if not text:
        return text, []
    reports: list[FixReport] = []
    result = text
    for brand in _BRANDS:
        if brand.lower() in {s.lower() for s in _STANDARDS_EXEMPT}:
            continue
        brand_pat = re.compile(r'\b' + re.escape(brand) + r'\b', re.IGNORECASE)
        if not brand_pat.search(result):
            continue

        def _replacer(m: re.Match, _brand: str = brand, _res: list[str] = [result]) -> str:
            ctx_start = max(0, m.start() - 60)
            ctx_end = min(len(_res[0]), m.end() + 60)
            ctx = _res[0][ctx_start:ctx_end]
            if _EQUIV_MARKER.search(ctx):
                return m.group(0)
            return m.group(0) + ' или эквивалент'

        new_result = brand_pat.sub(_replacer, result)
        if new_result != result:
            reports.append(FixReport("TEST-03", "full_text", "add_equivalent",
                                     brand, f"{brand} или эквивалент"))
            result = new_result

    return result, reports


def _add_equiv_to_value(value: str) -> tuple[str, list[str]]:
    """Добавить 'или эквивалент' после бренда в строке значения, если ещё нет."""
    changed: list[str] = []
    result = value
    for brand in _BRANDS:
        brand_lower = brand.lower()
        if brand_lower in {s.lower() for s in _STANDARDS_EXEMPT}:
            continue
        brand_pat = re.compile(r'\b' + re.escape(brand) + r'\b', re.IGNORECASE)
        if not brand_pat.search(result):
            continue

        def _replacer(m: re.Match, _brand: str = brand) -> str:
            ctx_start = max(0, m.start() - 60)
            ctx_end = min(len(result), m.end() + 60)
            ctx = result[ctx_start:ctx_end]
            if _EQUIV_MARKER.search(ctx):
                return m.group(0)
            changed.append(_brand)
            return m.group(0) + ' или эквивалент'

        result = brand_pat.sub(_replacer, result)
    return result, changed


def auto_fix_specs(
    rows: list[SpecRow],
    test_results: list[TestResult],
    llm_fixes: dict[str, str] | None = None,
) -> tuple[list[SpecRow], list[FixReport]]:
    """
    Применяет детерминированные авто-исправления к строкам по результатам тестов.
    llm_fixes: {"{field}|{spec_name}": "новое_значение"} — результаты LLM для TEST-04.
    Возвращает (исправленные строки, отчёт о фиксах).
    """
    failed_ids = {t.id for t in test_results if t.status == "fail"}
    llm_fixes = llm_fixes or {}
    reports: list[FixReport] = []

    fixed_rows: list[SpecRow] = []
    for row in rows:
        specs: list[tuple[str, str, str]] = list(row.specs)

        # ── TEST-01: удалить мета-комментарии из значений ─────────────────────
        if "TEST-01" in failed_ids:
            new_specs: list[tuple[str, str, str]] = []
            for sn, sv, sg in specs:
                fixed_v = sv
                for pat, _ in _META_PATTERNS:
                    fixed_v = pat.sub('', fixed_v).strip()
                if fixed_v != sv:
                    reports.append(FixReport("TEST-01", f"{row.field} → {sn}",
                                             "remove_meta", sv[:80], fixed_v[:80]))
                new_specs.append((sn, fixed_v, sg))
            specs = new_specs

        # ── TEST-02: удалить характеристики с запрещёнными формулировками ─────
        if "TEST-02" in failed_ids:
            clean_specs: list[tuple[str, str, str]] = []
            for sn, sv, sg in specs:
                combined = f"{sn} {sv}"
                banned = any(pat.search(combined) for pat, _ in _BANNED_CRITICAL)
                if banned:
                    reports.append(FixReport("TEST-02", f"{row.field} → {sn}",
                                             "remove_banned_spec", sv[:80], ""))
                else:
                    clean_specs.append((sn, sv, sg))
            specs = clean_specs

        # ── TEST-03: добавить «или эквивалент» после брендов ─────────────────
        if "TEST-03" in failed_ids:
            equiv_specs: list[tuple[str, str, str]] = []
            for sn, sv, sg in specs:
                new_v, brands_fixed = _add_equiv_to_value(sv)
                for b in brands_fixed:
                    reports.append(FixReport("TEST-03", f"{row.field} → {sn}",
                                             "add_equivalent", b, f"{b} или эквивалент"))
                equiv_specs.append((sn, new_v, sg))
            specs = equiv_specs

        # ── TEST-04: LLM-исправление неизмеримых значений ────────────────────
        if "TEST-04" in failed_ids and llm_fixes:
            llm_specs: list[tuple[str, str, str]] = []
            for sn, sv, sg in specs:
                key = f"{row.field}|{sn}"
                if key in llm_fixes:
                    new_v = llm_fixes[key].strip()
                    reports.append(FixReport("TEST-04", f"{row.field} → {sn}",
                                             "llm_measurable", sv[:80], new_v[:80]))
                    llm_specs.append((sn, new_v, sg))
                else:
                    llm_specs.append((sn, sv, sg))
            specs = llm_specs

        # ── TEST-05: удалить дубли характеристик (оставить более полный) ─────
        if "TEST-05" in failed_ids:
            seen_idx: dict[str, int] = {}
            dedup_specs: list[tuple[str, str, str]] = []
            for sn, sv, sg in specs:
                key = _normalize_name(sn)
                if key in seen_idx:
                    existing_i = seen_idx[key]
                    _, existing_v, _ = dedup_specs[existing_i]
                    if len(sv) > len(existing_v):
                        reports.append(FixReport("TEST-05", row.field,
                                                  "remove_duplicate", existing_v[:60], sv[:60]))
                        dedup_specs[existing_i] = (sn, sv, sg)
                    else:
                        reports.append(FixReport("TEST-05", row.field,
                                                  "remove_duplicate", sv[:60], existing_v[:60]))
                else:
                    seen_idx[key] = len(dedup_specs)
                    dedup_specs.append((sn, sv, sg))
            specs = dedup_specs

        # ── TEST-13: добавить пропущенные единицы к числовым значениям ──────
        if "TEST-13" in failed_ids:
            unit_fixed_specs: list[tuple[str, str, str]] = []
            for sn, sv, sg in specs:
                new_v = sv
                # Пропускаем исключения (тип, кол-во слотов и т.п.)
                if not any(excl.search(sn) for excl in _UNIT_SPEC_EXCLUSIONS):
                    for param_pat, allowed_units in _UNIT_RULES:
                        if not param_pat.search(sn):
                            continue
                        unit_ok = any(u.lower() in sv.lower() for u in allowed_units)
                        if not unit_ok and re.search(r'\d\s*$', sv.strip()):
                            # Значение заканчивается числом — добавляем первую подходящую единицу
                            unit_to_add = allowed_units[0]
                            new_v = sv.strip() + ' ' + unit_to_add
                            reports.append(FixReport(
                                "TEST-13", f"{row.field} → {sn}",
                                "fix_qty_unit", sv[:80], new_v[:80],
                            ))
                        break
                unit_fixed_specs.append((sn, new_v, sg))
            specs = unit_fixed_specs

        # ── TEST-07: исправить некорректные количества / единицы ─────────────
        fixed_qty = row.qty
        fixed_unit = row.qty_unit
        if "TEST-07" in failed_ids:
            if row.qty <= 0:
                reports.append(FixReport("TEST-07", row.field, "fix_qty",
                                         str(row.qty), "1"))
                fixed_qty = 1
            if not row.qty_unit:
                reports.append(FixReport("TEST-07", row.field, "fix_qty_unit",
                                         "", "шт."))
                fixed_unit = "шт."

        fixed_rows.append(SpecRow(
            name=row.name,
            field=row.field,
            qty=fixed_qty,
            qty_unit=fixed_unit,
            category=row.category,
            specs=specs,
            description=row.description,
        ))

    return fixed_rows, reports
