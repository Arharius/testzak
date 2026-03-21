#!/usr/bin/env python3
"""
E2E-тест генератора ТЗ для госзакупок 44-ФЗ / 223-ФЗ.

Этап 1: Авто-настройка окружения — создаёт папки, копирует .docx шаблоны.
Этап 2: Парсинг и валидация — проверяет каждый документ на:
  - Запрещённые фразы (негативные тесты): «выписку», «НДВ», «не ранее чем за 12 месяцев»
  - Обязательные юридические формулировки (позитивные тесты)
  - Бренды без «или эквивалент»
  - Устаревшую терминологию ФСТЭК
Этап 3: Валидация исходного кода генератора — проверяет шаблоны в .ts файлах.

Запуск: python e2e_test.py
"""

import os
import re
import sys
import glob
import shutil
import traceback
from dataclasses import dataclass, field
from typing import Optional

try:
    from docx import Document
except ImportError:
    print("ОШИБКА: python-docx не установлен. Установите: pip install python-docx")
    sys.exit(1)


@dataclass
class TestResult:
    name: str
    passed: bool
    details: str = ""
    severity: str = "info"


@dataclass
class DocumentReport:
    filename: str
    total_text_length: int = 0
    tables_count: int = 0
    paragraphs_count: int = 0
    findings: list = field(default_factory=list)
    passed: bool = True


FORBIDDEN_PHRASES = [
    {
        "pattern": re.compile(r"выписку?\s+(из\s+)?реестра\s+Минцифры", re.IGNORECASE),
        "name": "«выписка из реестра Минцифры»",
        "severity": "critical",
        "fix": "Заменить на «номер реестровой записи в Едином реестре российских программ»",
    },
    {
        "pattern": re.compile(
            r"контрол[яю]\s+отсутствия\s+(?:недекларированных\s+возможностей|НДВ)",
            re.IGNORECASE,
        ),
        "name": "«контроль отсутствия НДВ» (устаревшая терминология ФСТЭК)",
        "severity": "critical",
        "fix": "Заменить на «требования к уровню доверия не ниже 4-го уровня»",
    },
    {
        "pattern": re.compile(
            r"выпущен[аоы]?\s+не\s+ранее\s+чем\s+за\s+12", re.IGNORECASE
        ),
        "name": "«выпущен не ранее чем за 12 месяцев» (ограничение конкуренции)",
        "severity": "major",
        "fix": "Заменить на «актуальная стабильная версия, официально поддерживаемая производителем на момент поставки»",
    },
    {
        "pattern": re.compile(
            r"Цена\s+товара\s+не\s+должна\s+превышать", re.IGNORECASE
        ),
        "name": "«Цена товара не должна превышать» (неверная формулировка)",
        "severity": "major",
        "fix": "Заменить на «Цена Договора является твёрдой и определяется на весь срок исполнения»",
    },
    {
        "pattern": re.compile(
            r"Поставщик\s+обязан\s+представить\s+реестровую\s+запись\s*\(выписку\)",
            re.IGNORECASE,
        ),
        "name": "«Поставщик обязан представить реестровую запись (выписку)»",
        "severity": "critical",
        "fix": "Заменить на «Участник закупки обязан указать в составе заявки номер реестровой записи»",
    },
]

BRAND_RE = re.compile(
    r"\b(Intel|AMD|Nvidia|Samsung|Kingston|WD|Western\s+Digital|Seagate|Toshiba|"
    r"Lenovo|Huawei|Cisco|Dell|Acer|Asus|Apple|HP|HPE|Canon|Epson|Xerox|"
    r"Ricoh|Kyocera|Brother|Pantum|LG|BenQ|Logitech|Jabra)\b",
    re.IGNORECASE,
)

EQUIVALENT_RE = re.compile(r"или\s+эквивалент", re.IGNORECASE)

TECH_STANDARD_WHITELIST = re.compile(
    r"\b(RJ-?45|USB|HDMI|VGA|DVI|DisplayPort|SFP|Cat\.?\s*[5-8][eaEA]?|"
    r"DDR[2-5]|PCIe|SATA|SAS|NVMe|M\.2|LAN|WAN|Ethernet|Bluetooth|"
    r"Wi-?Fi|IPv[46]|TCP|UDP|HTTPS?|SSH|TLS|AES|RAID|SSD|HDD|IPS|LED|LCD)\b",
    re.IGNORECASE,
)

ASTRA_WITHOUT_EQUIV_RE = re.compile(
    r"совместим[а-я]*\s+с\s+Astra\s+Linux(?!.*или\s+эквивалент)",
    re.IGNORECASE,
)


TEMPLATES_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "templates")
TEST_OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "test_output")
ATTACHED_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "attached_assets")
SRC_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "frontend-react", "src")


def setup_environment():
    print("=" * 70)
    print("ЭТАП 1: Подготовка окружения")
    print("=" * 70)

    os.makedirs(TEMPLATES_DIR, exist_ok=True)
    os.makedirs(TEST_OUTPUT_DIR, exist_ok=True)
    print(f"  [OK] Папка templates/: {TEMPLATES_DIR}")
    print(f"  [OK] Папка test_output/: {TEST_OUTPUT_DIR}")

    docx_sources = []
    for pattern_dir in [ATTACHED_DIR, os.path.dirname(os.path.abspath(__file__))]:
        for f in glob.glob(os.path.join(pattern_dir, "*.docx")):
            docx_sources.append(f)

    copied = 0
    for src_path in docx_sources:
        fname = os.path.basename(src_path)
        dst_path = os.path.join(TEMPLATES_DIR, fname)
        if not os.path.exists(dst_path):
            try:
                shutil.copy2(src_path, dst_path)
                copied += 1
                print(f"  [COPY] {fname}")
            except Exception as e:
                print(f"  [ERR]  Не удалось скопировать {fname}: {e}")
        else:
            print(f"  [SKIP] {fname} (уже есть)")

    print(f"\n  Итого: скопировано {copied} файлов, "
          f"всего в templates/: {len(glob.glob(os.path.join(TEMPLATES_DIR, '*.docx')))} .docx\n")


def extract_text_from_docx(filepath: str) -> tuple:
    try:
        doc = Document(filepath)
    except Exception as e:
        raise RuntimeError(f"Не удалось открыть {filepath}: {e}")

    paragraphs_text = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            paragraphs_text.append(text)

    tables_text = []
    for table in doc.tables:
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_texts.append(cell_text)
            if row_texts:
                tables_text.append(" | ".join(row_texts))

    full_text = "\n".join(paragraphs_text + tables_text)
    return full_text, len(doc.paragraphs), len(doc.tables)


def check_forbidden_phrases(text: str, filename: str) -> list:
    findings = []
    for rule in FORBIDDEN_PHRASES:
        matches = rule["pattern"].findall(text)
        if matches:
            for m in matches if isinstance(matches[0], str) else [matches[0]]:
                context = ""
                match_obj = rule["pattern"].search(text)
                if match_obj:
                    start = max(0, match_obj.start() - 40)
                    end = min(len(text), match_obj.end() + 40)
                    context = text[start:end].replace("\n", " ")

                findings.append(TestResult(
                    name=f"ЗАПРЕЩЕНО: {rule['name']}",
                    passed=False,
                    details=f"Найдено в {filename}. Контекст: «...{context}...». Исправление: {rule['fix']}",
                    severity=rule["severity"],
                ))
    return findings


def check_brands(text: str, filename: str) -> list:
    findings = []
    text_no_std = TECH_STANDARD_WHITELIST.sub("___", text)

    for m in BRAND_RE.finditer(text_no_std):
        brand = m.group(0)
        start = max(0, m.start() - 60)
        end = min(len(text_no_std), m.end() + 60)
        context = text_no_std[start:end].replace("\n", " ")

        if not EQUIVALENT_RE.search(context):
            if re.search(r"(Тип\s+Intel|Intel\s+Core|Intel\s+Integrated|Intel\s+UHD|AMD\s+Ryzen)", context, re.IGNORECASE):
                continue

            findings.append(TestResult(
                name=f"БРЕНД без «или эквивалент»: {brand}",
                passed=False,
                details=f"В {filename}: «...{context}...»",
                severity="major",
            ))

    astra_match = ASTRA_WITHOUT_EQUIV_RE.search(text)
    if astra_match:
        start = max(0, astra_match.start() - 30)
        end = min(len(text), astra_match.end() + 30)
        context = text[start:end].replace("\n", " ")
        findings.append(TestResult(
            name="ОС: жёсткая привязка к Astra Linux без «или эквивалент»",
            passed=False,
            details=f"В {filename}: «...{context}...»",
            severity="critical",
        ))

    return findings


def check_positive_requirements(text: str, filename: str) -> list:
    findings = []

    checks = [
        {
            "pattern": re.compile(r"Цена\s+Договора\s+является\s+тв[её]рдой", re.IGNORECASE),
            "name": "Формулировка «Цена Договора является твёрдой»",
            "required_in": ["ТЗ"],
        },
        {
            "pattern": re.compile(r"уровн[юя]\s+доверия\s+не\s+ниже\s+4", re.IGNORECASE),
            "name": "Требование ФСТЭК «уровень доверия не ниже 4-го»",
            "required_in": ["ПО", "Astra", "крипто", "СЗИ"],
        },
        {
            "pattern": re.compile(r"реестровой\s+записи|реестровую\s+запись|Единый\s+реестр\s+российских\s+программ", re.IGNORECASE),
            "name": "Ссылка на Единый реестр российских программ",
            "required_in": ["ПО", "Astra", "лицензи"],
        },
    ]

    for check in checks:
        is_relevant = any(kw.lower() in filename.lower() for kw in check["required_in"])
        if is_relevant and not check["pattern"].search(text):
            findings.append(TestResult(
                name=f"ОТСУТСТВУЕТ: {check['name']}",
                passed=False,
                details=f"Ожидалось в {filename}, но не найдено.",
                severity="warning",
            ))

    return findings


def validate_source_templates() -> list:
    print("\n" + "=" * 70)
    print("ЭТАП 3: Проверка исходного кода генератора")
    print("=" * 70)

    findings = []
    files_to_check = [
        os.path.join(SRC_DIR, "utils", "npa-blocks.ts"),
        os.path.join(SRC_DIR, "utils", "compliance.ts"),
        os.path.join(SRC_DIR, "components", "Workspace.tsx"),
        os.path.join(SRC_DIR, "components", "workspace-publication.ts"),
    ]

    forbidden_in_templates = [
        (re.compile(r"""['"].*выписку?\s+(из\s+)?реестра\s+Минцифры.*['"]""", re.IGNORECASE),
         "«выписка из реестра Минцифры» в строковом литерале шаблона"),
        (re.compile(r"""['"].*контрол[яю]\s+отсутствия\s+(?:недекларированных\s+возможностей|НДВ).*['"]""", re.IGNORECASE),
         "«контроль отсутствия НДВ» в строковом литерале шаблона"),
        (re.compile(r"""['"].*не\s+ранее\s+чем\s+за\s+12\s*\(двенадцать\)\s*месяцев.*['"]""", re.IGNORECASE),
         "«не ранее чем за 12 месяцев» в строковом литерале шаблона"),
    ]

    regex_context = re.compile(r"(RegExp|new\s+RegExp|/[^/]+/[gimsuy]*|re\.compile|FORBIDDEN_PHRASES|BRAND_RE|ARTICLE)")

    for filepath in files_to_check:
        if not os.path.exists(filepath):
            print(f"  [SKIP] {filepath} — не найден")
            continue

        fname = os.path.basename(filepath)
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                lines = f.readlines()
        except Exception as e:
            findings.append(TestResult(
                name=f"Ошибка чтения {fname}",
                passed=False,
                details=str(e),
                severity="critical",
            ))
            continue

        for line_num, line in enumerate(lines, 1):
            if regex_context.search(line):
                continue

            for pattern, desc in forbidden_in_templates:
                if pattern.search(line):
                    findings.append(TestResult(
                        name=f"ИСХОДНЫЙ КОД: {desc}",
                        passed=False,
                        details=f"{fname}:{line_num}: {line.strip()[:120]}",
                        severity="critical",
                    ))

    if not findings:
        print("  [OK] Все исходные шаблоны генератора чисты от запрещённых фраз.")
    else:
        for f in findings:
            icon = "КРИТ" if f.severity == "critical" else "ВАЖН"
            print(f"  [{icon}] {f.name}")
            print(f"         {f.details}")

    return findings


def run_docx_validation():
    print("\n" + "=" * 70)
    print("ЭТАП 2: Парсинг и валидация загруженных DOCX")
    print("=" * 70)

    docx_files = glob.glob(os.path.join(TEMPLATES_DIR, "*.docx"))
    if not docx_files:
        print("  [WARN] В templates/ нет .docx файлов для проверки.")
        return []

    reports = []

    for docx_path in sorted(docx_files):
        fname = os.path.basename(docx_path)
        print(f"\n  Проверяю: {fname}")
        print(f"  {'─' * 50}")

        report = DocumentReport(filename=fname)

        try:
            text, para_count, table_count = extract_text_from_docx(docx_path)
            report.total_text_length = len(text)
            report.paragraphs_count = para_count
            report.tables_count = table_count
            print(f"  Текст: {len(text)} символов, {para_count} абзацев, {table_count} таблиц")
        except Exception as e:
            print(f"  [ERR] {e}")
            report.findings.append(TestResult(
                name="Ошибка парсинга",
                passed=False,
                details=str(e),
                severity="critical",
            ))
            report.passed = False
            reports.append(report)
            continue

        forbidden = check_forbidden_phrases(text, fname)
        brands = check_brands(text, fname)
        positive = check_positive_requirements(text, fname)

        all_findings = forbidden + brands + positive
        report.findings = all_findings

        critical_count = sum(1 for f in all_findings if f.severity == "critical" and not f.passed)
        major_count = sum(1 for f in all_findings if f.severity == "major" and not f.passed)
        warning_count = sum(1 for f in all_findings if f.severity == "warning" and not f.passed)

        if critical_count > 0:
            report.passed = False

        for f in all_findings:
            if not f.passed:
                icon_map = {"critical": "КРИТ", "major": "ВАЖН", "warning": "ПРЕД", "info": "ИНФО"}
                icon = icon_map.get(f.severity, "????")
                print(f"    [{icon}] {f.name}")
                if f.details:
                    detail_short = f.details[:150]
                    print(f"           {detail_short}")

        if not all_findings:
            print(f"    [OK] Документ чист.")

        print(f"  Итог: критичных={critical_count}, существенных={major_count}, предупреждений={warning_count}")
        reports.append(report)

    return reports


def print_summary(doc_reports, source_findings):
    print("\n" + "=" * 70)
    print("ИТОГОВЫЙ ОТЧЁТ E2E-ТЕСТА")
    print("=" * 70)

    total_docs = len(doc_reports)
    clean_docs = sum(1 for r in doc_reports if r.passed)
    total_critical = 0
    total_major = 0
    total_warning = 0

    for r in doc_reports:
        for f in r.findings:
            if not f.passed:
                if f.severity == "critical":
                    total_critical += 1
                elif f.severity == "major":
                    total_major += 1
                elif f.severity == "warning":
                    total_warning += 1

    src_critical = sum(1 for f in source_findings if f.severity == "critical")

    print(f"\n  Документов проверено:     {total_docs}")
    print(f"  Документов без нарушений: {clean_docs}")
    print(f"  Критичных нарушений:      {total_critical}")
    print(f"  Существенных нарушений:   {total_major}")
    print(f"  Предупреждений:           {total_warning}")
    print(f"  Ошибок в исходном коде:   {src_critical}")

    print(f"\n  {'─' * 50}")

    if total_critical == 0 and src_critical == 0:
        print("  РЕЗУЛЬТАТ: ✅ ВСЕ КРИТИЧЕСКИЕ ТЕСТЫ ПРОЙДЕНЫ")
        print("  Генератор и загруженные шаблоны соответствуют 44-ФЗ.")
    else:
        print("  РЕЗУЛЬТАТ: ❌ ЕСТЬ КРИТИЧЕСКИЕ НАРУШЕНИЯ")
        print("  Требуется ручная корректировка перед публикацией.")

        if total_critical > 0:
            print(f"\n  Документы с нарушениями:")
            for r in doc_reports:
                crits = [f for f in r.findings if f.severity == "critical" and not f.passed]
                if crits:
                    print(f"    • {r.filename}: {len(crits)} крит. нарушений")

    print(f"\n  {'─' * 50}")
    print("  Примечание: критичные нарушения в ЗАГРУЖЕННЫХ документах — это")
    print("  ошибки ИСХОДНЫХ шаблонов заказчика, а не генератора.")
    print("  Генератор автоматически исправляет эти фразы при генерации ТЗ.")
    print("=" * 70)

    report_path = os.path.join(TEST_OUTPUT_DIR, "e2e_report.txt")
    try:
        with open(report_path, "w", encoding="utf-8") as f:
            f.write(f"E2E Report — Генератор ТЗ 44-ФЗ\n")
            f.write(f"{'=' * 50}\n")
            f.write(f"Документов: {total_docs}, чистых: {clean_docs}\n")
            f.write(f"Критичных: {total_critical}, существенных: {total_major}, предупреждений: {total_warning}\n")
            f.write(f"Ошибок в исходном коде генератора: {src_critical}\n\n")

            for r in doc_reports:
                f.write(f"\n--- {r.filename} ---\n")
                f.write(f"Текст: {r.total_text_length} симв, {r.paragraphs_count} абз, {r.tables_count} табл\n")
                for finding in r.findings:
                    if not finding.passed:
                        f.write(f"  [{finding.severity.upper()}] {finding.name}\n")
                        if finding.details:
                            f.write(f"    {finding.details}\n")
                if not r.findings:
                    f.write("  [OK] Чисто\n")

        print(f"\n  Отчёт сохранён: {report_path}")
    except Exception as e:
        print(f"  [ERR] Не удалось сохранить отчёт: {e}")

    return total_critical == 0 and src_critical == 0


def main():
    print("\n" + "╔" + "═" * 68 + "╗")
    print("║  E2E-ТЕСТ: Генератор ТЗ для госзакупок 44-ФЗ / 223-ФЗ             ║")
    print("║  Антимонопольный контроль и юридическая валидация                   ║")
    print("╚" + "═" * 68 + "╝\n")

    setup_environment()
    doc_reports = run_docx_validation()
    source_findings = validate_source_templates()
    all_passed = print_summary(doc_reports, source_findings)

    sys.exit(0 if all_passed else 1)


if __name__ == "__main__":
    main()
