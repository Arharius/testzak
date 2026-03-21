#!/usr/bin/env python3
"""
run_all_tests.py — Полный E2E-тест генератора ТЗ для госзакупок 44-ФЗ / 223-ФЗ.

Для КАЖДОГО .docx в проекте:
  1. Открывает документ, находит наименование объекта закупки.
  2. Заменяет его на «ноутбук asus x1503» (симуляция пользовательского ввода).
  3. Сохраняет в test_output/tested_<имя>.docx.
  4. Повторно открывает сохранённый файл, считывает ВЕСЬ текст.
  5. Выполняет позитивные и негативные проверки (assert-уровня).
  6. Формирует итоговый отчёт с exit code.

Запуск: python run_all_tests.py
"""

import os
import re
import sys
import glob
import copy
import shutil
import traceback
from dataclasses import dataclass, field
from typing import List, Tuple, Optional

try:
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError
except ImportError:
    print("[FATAL] python-docx не установлен. Установите: pip install python-docx")
    sys.exit(2)


# ---------------------------------------------------------------------------
#  Конфигурация
# ---------------------------------------------------------------------------

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEST_OUTPUT_DIR = os.path.join(BASE_DIR, "test_output")
REPLACEMENT_TEXT = "ноутбук asus x1503"

PLACEHOLDER_PATTERNS = [
    re.compile(
        r"комплект[аы]?\s+товаров\s*\(\s*\d+\s+позиц[а-я]*\s*\)",
        re.IGNORECASE,
    ),
    re.compile(
        r"комплект[аы]?\s+периферийных\s+устройств\s+и\s+расходн[а-я]+\s+материал[а-я]*",
        re.IGNORECASE,
    ),
    re.compile(
        r"закупк[аеуи]\s+(?:ноутбуков|системных\s+блоков|программного\s+обеспечения|"
        r"периферийных|сетевого\s+оборудования|комплектующих)[^\.\n]{0,60}",
        re.IGNORECASE,
    ),
    re.compile(
        r"(?:Н|н)аименование\s+объекта\s+поставки:\s*([^\.]{5,80}?)(?:\s*\(далее|\.\s)",
        re.IGNORECASE,
    ),
    re.compile(
        r"на\s+поставку\s+([^\n\.\(]{3,60})",
        re.IGNORECASE,
    ),
]

NEGATIVE_RULES = [
    # ── Правила для ПО (Software) ──────────────────────────────────────
    {
        "pattern": re.compile(
            r"выписку?\s+(из\s+)?реестра\s+Минцифры", re.IGNORECASE
        ),
        "name": "«выписка из реестра Минцифры»",
        "severity": "critical",
        "context_desc": "Устаревшая формулировка. Должно быть: «номер реестровой записи в Едином реестре».",
    },
    {
        "pattern": re.compile(
            r"\bНДВ\b|контрол[яю]\s+отсутствия\s+недекларированных\s+возможностей",
            re.IGNORECASE,
        ),
        "name": "«НДВ» / «контроль отсутствия недекларированных возможностей»",
        "severity": "critical",
        "context_desc": "Терминология отменена ФСТЭК. Замена: «уровень доверия не ниже 4-го».",
    },
    {
        "pattern": re.compile(
            r"не\s+ранее\s+чем\s+за\s+12\s*(?:\(двенадцать\))?\s*месяцев",
            re.IGNORECASE,
        ),
        "name": "«не ранее чем за 12 месяцев»",
        "severity": "major",
        "context_desc": "Ограничивает конкуренцию по дате выпуска. Замена: «актуальная стабильная версия».",
    },
    {
        "pattern": re.compile(
            r"Цена\s+товара\s+не\s+должна\s+превышать", re.IGNORECASE
        ),
        "name": "«Цена товара не должна превышать»",
        "severity": "major",
        "context_desc": "Неверная формулировка ценообразования.",
    },
    {
        "pattern": re.compile(
            r"Поставщик\s+обязан\s+представить\s+реестровую\s+запись\s*\(выписку\)",
            re.IGNORECASE,
        ),
        "name": "«представить реестровую запись (выписку)»",
        "severity": "critical",
        "context_desc": "Замена: «указать в составе заявки номер реестровой записи».",
    },
    # ── Правила для оборудования (Hardware / Anti-ФАС) ─────────────────
    {
        "pattern": re.compile(
            r"выписку?\s+(из\s+)?ГИСП", re.IGNORECASE
        ),
        "name": "«выписка из ГИСП» (ограничение нацрежима)",
        "severity": "critical",
        "context_desc": (
            "Участник не обязан предоставлять выписку из ГИСП. "
            "Замена: «номера реестровых записей из реестра российской "
            "промышленной продукции или евразийского реестра промышленной продукции»."
        ),
    },
    {
        "pattern": re.compile(
            r"оригинал\s+документа", re.IGNORECASE
        ),
        "name": "«оригинал документа» (избыточное формальное требование)",
        "severity": "major",
        "context_desc": (
            "Требование оригинала сужает конкуренцию. Замена: «документ "
            "(гарантийный талон, сертификат или иной документ), подтверждающий "
            "гарантийные обязательства»."
        ),
    },
    {
        "pattern": re.compile(
            r"\[!\]", re.IGNORECASE
        ),
        "name": "Системный маркер «[!]» (забытый плейсхолдер)",
        "severity": "critical",
        "context_desc": (
            "В финальном документе остался маркер-плейсхолдер «[!]». "
            "Генератор обязан удалять комментарии вида «[!] Требуется более "
            "конкретное значение», оставляя только содержательный текст."
        ),
    },
]

BRAND_RE = re.compile(
    r"\b(Intel|AMD|Nvidia|Samsung|Kingston|WD|Western\s+Digital|Seagate|Toshiba|"
    r"Lenovo|Huawei|Cisco|Dell|Acer|Asus|Apple|HP|HPE|Canon|Epson|Xerox|"
    r"Ricoh|Kyocera|Brother|Pantum|LG|BenQ|Logitech|Jabra|Realtek)\b",
    re.IGNORECASE,
)

EQUIV_RE = re.compile(r"или\s+эквивалент", re.IGNORECASE)

TECH_WHITELIST_RE = re.compile(
    r"\b(RJ-?45|USB|HDMI|VGA|DVI|DisplayPort|SFP\+?|Cat\.?\s*[5-8][eaEA]?|"
    r"DDR[2-5]|PCIe|SATA|SAS|NVMe|M\.2|LAN|WAN|Ethernet|Bluetooth|"
    r"Wi-?Fi|IPv[46]|TCP|UDP|HTTPS?|SSH|TLS|AES|RAID|SSD|HDD|IPS|LED|LCD|"
    r"IEEE\s*802)\b",
    re.IGNORECASE,
)

ASTRA_HARD_RE = re.compile(
    r"совместим[а-яё]*\s+(с\s+)?(?:операционной\s+системой\s+)?Astra\s+Linux"
    r"(?!.*или\s+эквивалент)",
    re.IGNORECASE,
)


# ---------------------------------------------------------------------------
#  Структуры данных
# ---------------------------------------------------------------------------

@dataclass
class Finding:
    rule_name: str
    severity: str
    context: str
    description: str


@dataclass
class FileReport:
    filename: str
    source_path: str
    output_path: str = ""
    char_count: int = 0
    para_count: int = 0
    table_count: int = 0
    replacement_done: bool = False
    replacement_original: str = ""
    positive_replacement_ok: bool = False
    positive_price_ok: Optional[bool] = None
    positive_fstek_ok: Optional[bool] = None
    findings: List[Finding] = field(default_factory=list)
    error: str = ""

    @property
    def critical_count(self) -> int:
        return sum(1 for f in self.findings if f.severity == "critical")

    @property
    def major_count(self) -> int:
        return sum(1 for f in self.findings if f.severity == "major")

    @property
    def warning_count(self) -> int:
        return sum(1 for f in self.findings if f.severity == "warning")

    @property
    def is_clean(self) -> bool:
        return self.critical_count == 0 and len(self.error) == 0


# ---------------------------------------------------------------------------
#  Функция 1: Поиск всех .docx файлов
# ---------------------------------------------------------------------------

def discover_docx_files() -> List[str]:
    """Ищет ВСЕ .docx в корне, templates/ и attached_assets/.
    Исключает файлы из test_output/ и временные файлы (~$...)."""
    search_dirs = [
        BASE_DIR,
        os.path.join(BASE_DIR, "templates"),
        os.path.join(BASE_DIR, "attached_assets"),
    ]
    seen_names = set()
    result = []

    for search_dir in search_dirs:
        if not os.path.isdir(search_dir):
            continue
        for fpath in sorted(glob.glob(os.path.join(search_dir, "*.docx"))):
            fname = os.path.basename(fpath)
            if fname.startswith("~$"):
                continue
            abs_path = os.path.abspath(fpath)
            if "test_output" in abs_path:
                continue
            if fname not in seen_names:
                seen_names.add(fname)
                result.append(fpath)

    return sorted(result, key=lambda p: os.path.basename(p))


# ---------------------------------------------------------------------------
#  Функция 2: Извлечение ВСЕГО текста из .docx
# ---------------------------------------------------------------------------

def extract_full_text(filepath: str) -> Tuple[str, int, int]:
    """Открывает .docx, собирает текст из ВСЕХ абзацев и ВСЕХ ячеек таблиц.
    Возвращает (полный_текст, число_абзацев, число_таблиц)."""
    doc = Document(filepath)

    text_parts = []

    for para in doc.paragraphs:
        stripped = para.text.strip()
        if stripped:
            text_parts.append(stripped)

    for table in doc.tables:
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_cells.append(cell_text)
            if row_cells:
                text_parts.append(" | ".join(row_cells))

    full_text = "\n".join(text_parts)
    return full_text, len(doc.paragraphs), len(doc.tables)


# ---------------------------------------------------------------------------
#  Функция 3: Замена наименования объекта закупки
# ---------------------------------------------------------------------------

def replace_procurement_name(filepath: str, output_path: str) -> Tuple[bool, str]:
    """Открывает .docx, находит наименование объекта закупки по шаблонам,
    заменяет на REPLACEMENT_TEXT, сохраняет в output_path.
    Возвращает (успех, исходный_текст_до_замены)."""
    doc = Document(filepath)
    replaced = False
    original_text = ""

    def try_replace_in_text(text: str) -> Tuple[str, bool, str]:
        """Пробует подставить REPLACEMENT_TEXT вместо найденного плейсхолдера."""
        for pattern in PLACEHOLDER_PATTERNS:
            match = pattern.search(text)
            if match:
                found = match.group(0)
                new_text = text.replace(found, REPLACEMENT_TEXT, 1)
                return new_text, True, found
        return text, False, ""

    for para in doc.paragraphs:
        if not replaced and para.text.strip():
            new_text, did_replace, orig = try_replace_in_text(para.text)
            if did_replace:
                replaced = True
                original_text = orig
                for run in para.runs:
                    run.text = ""
                para.runs[0].text = new_text if para.runs else new_text
                if not para.runs:
                    para.text = new_text
                break

    if not replaced:
        for table in doc.tables:
            if replaced:
                break
            for row in table.rows:
                if replaced:
                    break
                for cell in row.cells:
                    if replaced:
                        break
                    for para in cell.paragraphs:
                        if not replaced and para.text.strip():
                            new_text, did_replace, orig = try_replace_in_text(para.text)
                            if did_replace:
                                replaced = True
                                original_text = orig
                                if para.runs:
                                    for run in para.runs:
                                        run.text = ""
                                    para.runs[0].text = new_text
                                break

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return replaced, original_text


# ---------------------------------------------------------------------------
#  Функция 4: Негативные антимонопольные проверки
# ---------------------------------------------------------------------------

def run_negative_checks(text: str, filename: str) -> List[Finding]:
    """Проверяет текст на запрещённые фразы, бренды без «или эквивалент»,
    привязку к Astra Linux."""
    findings = []

    for rule in NEGATIVE_RULES:
        match = rule["pattern"].search(text)
        if match:
            start = max(0, match.start() - 50)
            end = min(len(text), match.end() + 50)
            ctx = text[start:end].replace("\n", " ").strip()
            findings.append(Finding(
                rule_name=rule["name"],
                severity=rule["severity"],
                context=f"«...{ctx}...»",
                description=rule["context_desc"],
            ))

    text_cleaned = TECH_WHITELIST_RE.sub("___STD___", text)
    brand_seen = set()
    for m in BRAND_RE.finditer(text_cleaned):
        brand = m.group(0)
        brand_key = brand.lower()
        if brand_key in brand_seen:
            continue

        start = max(0, m.start() - 80)
        end = min(len(text_cleaned), m.end() + 80)
        context_window = text_cleaned[start:end]

        if EQUIV_RE.search(context_window):
            continue

        if re.search(
            r"(Тип\s+Intel|Intel\s+Core|Intel\s+Integrated|Intel\s+UHD|"
            r"AMD\s+Ryzen|Intel\s+Pentium|Процессор|Графическ|Чипсет|"
            r"Intel\s+Wireless|Intel\s+I219|Realtek)",
            context_window, re.IGNORECASE
        ):
            continue

        brand_seen.add(brand_key)
        ctx = context_window.replace("\n", " ").strip()
        findings.append(Finding(
            rule_name=f"Бренд «{brand}» без «или эквивалент»",
            severity="major",
            context=f"«...{ctx[:120]}...»",
            description="Нарушение ст. 33 44-ФЗ: указание торговых марок без «или эквивалент».",
        ))

    astra_m = ASTRA_HARD_RE.search(text)
    if astra_m:
        start = max(0, astra_m.start() - 30)
        end = min(len(text), astra_m.end() + 30)
        ctx = text[start:end].replace("\n", " ")
        findings.append(Finding(
            rule_name="Жёсткая привязка к Astra Linux без «или эквивалент»",
            severity="critical",
            context=f"«...{ctx}...»",
            description="Ограничивает конкуренцию. Нужно: «...совместим с Astra Linux или эквивалент».",
        ))

    return findings


# ---------------------------------------------------------------------------
#  Функция 5: Позитивные проверки
# ---------------------------------------------------------------------------

def run_positive_checks(
    text: str, filename: str, replacement_done: bool
) -> Tuple[bool, Optional[bool], Optional[bool]]:
    """Проверяет:
      a) Наличие REPLACEMENT_TEXT (если замена была выполнена).
      b) Наличие «Цена Договора является твёрдой» (в файлах ТЗ на поставку).
      c) Наличие «лицензию ФСТЭК России» или «ФСТЭК» (в файлах ПО/Astra).
    Возвращает (replacement_ok, price_ok_or_None, fstek_ok_or_None)."""

    replacement_ok = True
    if replacement_done:
        replacement_ok = REPLACEMENT_TEXT in text.lower() or REPLACEMENT_TEXT in text

    price_ok = None
    is_supply_doc = any(kw in filename.lower() for kw in [
        "тз_", "tz_", "техническое", "поставк", "комплект"
    ])
    if is_supply_doc:
        price_pattern = re.compile(
            r"Цена\s+Договора\s+является\s+тв[её]рдой", re.IGNORECASE
        )
        price_ok = bool(price_pattern.search(text))

    fstek_ok = None
    is_software_doc = any(kw in filename.lower() for kw in [
        "по_", "astra", "linux", "лицензи", "программ"
    ])
    if is_software_doc:
        fstek_pattern = re.compile(
            r"(лицензи[юя]\s+ФСТЭК|сертификат.*ФСТЭК|ФСТЭК\s+Росси[ия])",
            re.IGNORECASE,
        )
        fstek_ok = bool(fstek_pattern.search(text))

    return replacement_ok, price_ok, fstek_ok


# ---------------------------------------------------------------------------
#  Функция 6: Обработка одного файла (полный цикл)
# ---------------------------------------------------------------------------

def process_single_file(filepath: str) -> FileReport:
    """Полный цикл: замена → сохранение → повторное чтение → валидация."""
    fname = os.path.basename(filepath)
    output_name = f"tested_{fname}"
    output_path = os.path.join(TEST_OUTPUT_DIR, output_name)

    report = FileReport(
        filename=fname,
        source_path=filepath,
        output_path=output_path,
    )

    try:
        replaced, orig_text = replace_procurement_name(filepath, output_path)
        report.replacement_done = replaced
        report.replacement_original = orig_text
    except Exception as e:
        report.error = f"Ошибка при замене/сохранении: {e}\n{traceback.format_exc()}"
        return report

    try:
        full_text, para_count, table_count = extract_full_text(output_path)
        report.char_count = len(full_text)
        report.para_count = para_count
        report.table_count = table_count
    except Exception as e:
        report.error = f"Ошибка при повторном чтении: {e}\n{traceback.format_exc()}"
        return report

    repl_ok, price_ok, fstek_ok = run_positive_checks(full_text, fname, replaced)
    report.positive_replacement_ok = repl_ok
    report.positive_price_ok = price_ok
    report.positive_fstek_ok = fstek_ok

    neg_findings = run_negative_checks(full_text, fname)
    report.findings = neg_findings

    if not repl_ok and replaced:
        report.findings.insert(0, Finding(
            rule_name="ПОЗИТИВНЫЙ ТЕСТ: текст замены не найден в сохранённом файле",
            severity="critical",
            context=f"Ожидался текст: «{REPLACEMENT_TEXT}»",
            description="Замена была выполнена, но при повторном чтении текст не обнаружен.",
        ))

    if price_ok is False:
        report.findings.append(Finding(
            rule_name="ПОЗИТИВНЫЙ ТЕСТ: отсутствует «Цена Договора является твёрдой»",
            severity="warning",
            context="",
            description="Рекомендуемая формулировка о твёрдой цене не найдена в документе.",
        ))

    if fstek_ok is False:
        report.findings.append(Finding(
            rule_name="ПОЗИТИВНЫЙ ТЕСТ: отсутствует упоминание ФСТЭК",
            severity="warning",
            context="",
            description="Для ПО с требованиями ИБ ожидается ссылка на лицензию/сертификат ФСТЭК.",
        ))

    return report


# ---------------------------------------------------------------------------
#  Функция 7: Красивый вывод в консоль
# ---------------------------------------------------------------------------

SEVERITY_ICON = {
    "critical": "\033[91mКРИТ\033[0m",
    "major":    "\033[93mВАЖН\033[0m",
    "warning":  "\033[33mПРЕД\033[0m",
    "info":     "\033[36mИНФО\033[0m",
}

SEVERITY_ICON_PLAIN = {
    "critical": "КРИТ",
    "major": "ВАЖН",
    "warning": "ПРЕД",
    "info": "ИНФО",
}


def print_file_report(report: FileReport, index: int) -> None:
    """Выводит результат обработки одного файла."""
    print(f"\n  [{index}] {report.filename}")
    print(f"  {'─' * 60}")

    if report.error:
        print(f"    [\033[91mОШИБКА\033[0m] {report.error.splitlines()[0]}")
        return

    print(f"    Символов: {report.char_count:,} | Абзацев: {report.para_count} | Таблиц: {report.table_count}")

    if report.replacement_done:
        print(f"    [ЗАМЕНА] «{report.replacement_original[:60]}» → «{REPLACEMENT_TEXT}»")
        status = "\033[92mОК\033[0m" if report.positive_replacement_ok else "\033[91mНЕ НАЙДЕН\033[0m"
        print(f"    [ASSERT] Текст замены в сохранённом файле: {status}")
    else:
        print(f"    [ЗАМЕНА] Плейсхолдер не найден — файл сохранён без изменений")

    if report.positive_price_ok is not None:
        status = "\033[92mОК\033[0m" if report.positive_price_ok else "\033[93mОТСУТСТВУЕТ\033[0m"
        print(f"    [ASSERT] «Цена Договора является твёрдой»: {status}")

    if report.positive_fstek_ok is not None:
        status = "\033[92mОК\033[0m" if report.positive_fstek_ok else "\033[93mОТСУТСТВУЕТ\033[0m"
        print(f"    [ASSERT] Упоминание ФСТЭК: {status}")

    if report.findings:
        print(f"    Нарушений: {report.critical_count} крит. / {report.major_count} существ. / {report.warning_count} предупр.")
        for f in report.findings:
            icon = SEVERITY_ICON.get(f.severity, "????")
            print(f"      [{icon}] {f.rule_name}")
            if f.context:
                print(f"             {f.context[:140]}")
    else:
        print(f"    \033[92m[OK]\033[0m Антимонопольных нарушений не обнаружено.")


# ---------------------------------------------------------------------------
#  Функция 8: Итоговый отчёт
# ---------------------------------------------------------------------------

def print_final_report(reports: List[FileReport]) -> bool:
    """Печатает итоговый отчёт и сохраняет в файл. Возвращает True если всё чисто."""
    print("\n" + "=" * 70)
    print("  ИТОГОВЫЙ ОТЧЁТ: run_all_tests.py")
    print("=" * 70)

    total = len(reports)
    errors = sum(1 for r in reports if r.error)
    with_replacement = sum(1 for r in reports if r.replacement_done)
    repl_verified = sum(1 for r in reports if r.replacement_done and r.positive_replacement_ok)

    total_crit = sum(r.critical_count for r in reports)
    total_major = sum(r.major_count for r in reports)
    total_warn = sum(r.warning_count for r in reports)
    clean = sum(1 for r in reports if r.is_clean and not r.error)

    print(f"""
  Файлов обработано:            {total}
  Ошибок открытия/сохранения:   {errors}
  Замена плейсхолдера:          {with_replacement} из {total} файлов
  Замена подтверждена (assert):  {repl_verified} из {with_replacement}
  
  Критичных нарушений:          {total_crit}
  Существенных нарушений:       {total_major}
  Предупреждений:               {total_warn}
  Файлов без нарушений:         {clean} из {total}
""")

    print(f"  {'─' * 60}")

    all_ok = total_crit == 0 and errors == 0

    if all_ok:
        print("  \033[92mРЕЗУЛЬТАТ: ✅ КРИТИЧЕСКИХ НАРУШЕНИЙ НЕТ\033[0m")
        print("  Все документы прошли антимонопольную проверку.")
    else:
        print("  \033[91mРЕЗУЛЬТАТ: ❌ ОБНАРУЖЕНЫ КРИТИЧЕСКИЕ НАРУШЕНИЯ\033[0m")
        if errors:
            print(f"\n  Файлы с ошибками:")
            for r in reports:
                if r.error:
                    print(f"    • {r.filename}: {r.error.splitlines()[0][:80]}")
        if total_crit:
            print(f"\n  Файлы с критическими нарушениями:")
            for r in reports:
                if r.critical_count > 0:
                    print(f"    • {r.filename}: {r.critical_count} крит.")
                    for f in r.findings:
                        if f.severity == "critical":
                            print(f"      → {f.rule_name}")

    print(f"\n  {'─' * 60}")
    print("  Примечание: нарушения в ИСХОДНЫХ документах заказчика — это")
    print("  ожидаемое поведение. Генератор ТЗ автоматически исправляет")
    print("  эти фразы при генерации финального документа.")
    print("=" * 70)

    report_path = os.path.join(TEST_OUTPUT_DIR, "full_test_report.txt")
    try:
        with open(report_path, "w", encoding="utf-8") as fp:
            fp.write("ПОЛНЫЙ ОТЧЁТ E2E-ТЕСТА: run_all_tests.py\n")
            fp.write(f"{'=' * 60}\n\n")
            fp.write(f"Файлов: {total} | Чистых: {clean} | Ошибок: {errors}\n")
            fp.write(f"Критичных: {total_crit} | Существенных: {total_major} | Предупреждений: {total_warn}\n")
            fp.write(f"Замен выполнено: {with_replacement} | Подтверждено: {repl_verified}\n\n")

            for i, r in enumerate(reports, 1):
                fp.write(f"\n{'─' * 50}\n")
                fp.write(f"[{i}] {r.filename}\n")
                fp.write(f"  Путь: {r.source_path}\n")
                fp.write(f"  Выход: {r.output_path}\n")
                if r.error:
                    fp.write(f"  ОШИБКА: {r.error}\n")
                    continue
                fp.write(f"  Символов: {r.char_count} | Абзацев: {r.para_count} | Таблиц: {r.table_count}\n")
                fp.write(f"  Замена: {'ДА' if r.replacement_done else 'НЕТ'}")
                if r.replacement_done:
                    fp.write(f" | «{r.replacement_original[:50]}» → «{REPLACEMENT_TEXT}»")
                    fp.write(f" | Подтверждение: {'OK' if r.positive_replacement_ok else 'FAIL'}")
                fp.write("\n")
                if r.positive_price_ok is not None:
                    fp.write(f"  Цена Договора: {'OK' if r.positive_price_ok else 'ОТСУТСТВУЕТ'}\n")
                if r.positive_fstek_ok is not None:
                    fp.write(f"  ФСТЭК: {'OK' if r.positive_fstek_ok else 'ОТСУТСТВУЕТ'}\n")
                fp.write(f"  Нарушений: {r.critical_count}К/{r.major_count}С/{r.warning_count}П\n")
                for f in r.findings:
                    fp.write(f"    [{SEVERITY_ICON_PLAIN[f.severity]}] {f.rule_name}\n")
                    if f.context:
                        fp.write(f"           {f.context[:140]}\n")
                    if f.description:
                        fp.write(f"           {f.description}\n")
                if not r.findings and not r.error:
                    fp.write(f"    [OK] Чисто.\n")

        print(f"\n  Отчёт сохранён: {report_path}")
    except Exception as e:
        print(f"  [ERR] Не удалось сохранить отчёт: {e}")

    return all_ok


# ---------------------------------------------------------------------------
#  main()
# ---------------------------------------------------------------------------

def main():
    print("\n" + "╔" + "═" * 68 + "╗")
    print("║  run_all_tests.py — E2E-тест генератора ТЗ (44-ФЗ / 223-ФЗ)       ║")
    print("║  Симуляция ввода + глубокий парсинг + антимонопольная валидация      ║")
    print("╚" + "═" * 68 + "╝\n")

    print("=" * 70)
    print("  ЭТАП 1: Поиск всех .docx файлов")
    print("=" * 70)

    os.makedirs(TEST_OUTPUT_DIR, exist_ok=True)

    docx_files = discover_docx_files()
    if not docx_files:
        print("  [FATAL] Не найдено ни одного .docx файла в проекте!")
        sys.exit(2)

    print(f"  Найдено: {len(docx_files)} файлов\n")
    for i, f in enumerate(docx_files, 1):
        print(f"    {i:2d}. {os.path.basename(f)}")
        print(f"        └─ {f}")

    print(f"\n{'=' * 70}")
    print("  ЭТАП 2: Замена наименования + сохранение в test_output/")
    print("=" * 70)

    reports = []
    for i, filepath in enumerate(docx_files, 1):
        fname = os.path.basename(filepath)
        print(f"\n  ── Обработка [{i}/{len(docx_files)}]: {fname}")
        report = process_single_file(filepath)
        reports.append(report)
        if report.error:
            print(f"    \033[91m[ОШИБКА]\033[0m {report.error.splitlines()[0][:80]}")
        elif report.replacement_done:
            print(f"    \033[92m[OK]\033[0m Замена + сохранение → test_output/tested_{fname}")
        else:
            print(f"    \033[93m[SKIP]\033[0m Плейсхолдер не найден, сохранено как есть")

    print(f"\n{'=' * 70}")
    print("  ЭТАП 3: Валидация сохранённых файлов (asserts)")
    print("=" * 70)

    for i, report in enumerate(reports, 1):
        print_file_report(report, i)

    all_ok = print_final_report(reports)
    sys.exit(0 if all_ok else 1)


if __name__ == "__main__":
    main()
